"""
Plagiarism Checker Service
Ported from: https://github.com/cu-sanjay/Free-Turnitin-Plagiarism-Checker

This module provides plagiarism detection by:
1. Splitting text into sentences
2. Searching DuckDuckGo and CrossRef for each sentence
3. Fetching page content and calculating similarity
"""

import asyncio
import math
import random
import re
import time
import traceback
import xml.etree.ElementTree as ET
from dataclasses import dataclass
from datetime import datetime, timezone
from email.utils import parsedate_to_datetime
from io import BytesIO
from typing import Any, List, Optional, Protocol, Set
from urllib.parse import quote

import fitz  # PyMuPDF
import httpx
from bs4 import BeautifulSoup
from pydantic import BaseModel, ConfigDict, Field

import nltk
from nltk.stem import PorterStemmer

try:
    nltk.data.find("tokenizers/punkt")
except LookupError:
    nltk.download("punkt", quiet=True)

try:
    nltk.data.find("tokenizers/punkt_tab")
except LookupError:
    nltk.download("punkt_tab", quiet=True)

_stemmer = PorterStemmer()


def stem_word(word: str) -> str:
    return _stemmer.stem(word.lower())


def stem_text(text: str) -> Set[str]:
    words = re.sub(r"[^\w\s]", "", text.lower()).split()
    return {stem_word(w) for w in words if len(w) > 1}


from app.core.config import settings
from app.schemas.plagiarism import (
    MatchGroup,
    PlagiarismCheckRequest,
    PlagiarismCheckResponse,
    ReportV2,
    ReportV2Caveat,
    ReportV2SourceGroup,
    ReportV2SourceSpan,
    SentenceResult,
    SourceMatch,
)
from app.services.semantic_similarity import (
    calculate_semantic_similarity,
    get_quota_info,
)


# User agents for rotation to avoid blocking
USER_AGENTS = [
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:121.0) Gecko/20100101 Firefox/121.0",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/17.2 Safari/605.1.15",
]

# Semaphore to limit concurrent requests
MAX_CONCURRENT_REQUESTS = 5

ARXIV_API_URL = "https://export.arxiv.org/api/query"
ARXIV_TIMEOUT_SECONDS = 10.0
ARXIV_MAX_RETRIES = 3
ARXIV_BACKOFF_SECONDS = (2.0, 4.0, 8.0)
ARXIV_BACKOFF_JITTER_MAX_SECONDS = 0.5
ARXIV_POLITENESS_DELAY_SECONDS = 3.0
ARXIV_MAX_RESULTS_UPPER_BOUND = 25
ARXIV_MAX_START_WINDOW = 1000

CORE_API_URL = "https://api.core.ac.uk/v3/search/works"
CORE_TIMEOUT_SECONDS = 10.0
CORE_MAX_RETRIES = 3
CORE_BACKOFF_SECONDS = (2.0, 4.0, 8.0)
CORE_BACKOFF_JITTER_MAX_SECONDS = 0.5
CORE_MIN_REQUEST_INTERVAL_SECONDS = 2.0
CORE_CONNECTOR_CONCURRENCY = 1
CORE_DEGRADED_FAILURE_THRESHOLD = 3
CORE_DEGRADED_WINDOW_SECONDS = 60.0
CORE_AUTH_MODE_BEARER = "bearer"
CORE_AUTH_MODE_QUERY = "query"

PUBMED_ESEARCH_URL = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi"
PUBMED_ESUMMARY_URL = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esummary.fcgi"
PUBMED_MAX_RETRIES = 3
PUBMED_BACKOFF_SECONDS = (1.0, 2.0, 4.0)
PUBMED_BACKOFF_JITTER_MAX_SECONDS = 0.5
PUBMED_NO_KEY_MIN_INTERVAL_SECONDS = 0.4
PUBMED_KEY_MIN_INTERVAL_SECONDS = 0.12
PUBMED_CONNECTOR_CONCURRENCY = 2
PUBMED_RETMAX_UPPER_BOUND = 25
PUBMED_MAX_START_WINDOW = 100

SOURCE_CONNECTOR_PARALLELISM = 3
SIMILARITY_TOP_K_CANDIDATES = 5
SEARCH_RESULT_CACHE_MAX_ENTRIES = 64
MIN_SENTENCE_LENGTH_CHARS = 20
HEAVY_EXCLUSION_RATIO_THRESHOLD = 0.6
DUCKDUCKGO_BREAKER_FAILURE_THRESHOLD = 5
DUCKDUCKGO_BREAKER_COOLDOWN_SECONDS = 90.0
DUCKDUCKGO_BREAKER_OPEN_LOG_INTERVAL_SECONDS = 30.0
DUCKDUCKGO_ERROR_LOG_INTERVAL_SECONDS = 15.0
DUCKDUCKGO_CONNECT_TIMEOUT_SECONDS = 5.0
DUCKDUCKGO_CONNECT_MAX_RETRIES = 3
DUCKDUCKGO_BACKOFF_BASE_SECONDS = 1.0
DUCKDUCKGO_BACKOFF_JITTER_MAX_SECONDS = 0.4
DUCKDUCKGO_RETRYABLE_STATUS_CODES = {429, 500, 502, 503, 504}

_duckduckgo_failure_count = 0
_duckduckgo_open_until = 0.0
_duckduckgo_last_open_log_at = 0.0
_duckduckgo_last_error_log_at = 0.0

REFERENCE_SECTION_PATTERN = re.compile(
    r"(?im)^\s*(references|bibliography|works\s+cited|tai\s+lieu\s+tham\s+khao)\s*:?\s*$"
)


def _duckduckgo_breaker_is_open() -> bool:
    return time.monotonic() < _duckduckgo_open_until


def _duckduckgo_breaker_remaining_seconds() -> float:
    return max(0.0, _duckduckgo_open_until - time.monotonic())


def _duckduckgo_breaker_on_success() -> None:
    global _duckduckgo_failure_count, _duckduckgo_open_until
    _duckduckgo_failure_count = 0
    _duckduckgo_open_until = 0.0


def _duckduckgo_breaker_on_failure() -> None:
    global _duckduckgo_failure_count, _duckduckgo_open_until
    _duckduckgo_failure_count += 1
    if _duckduckgo_failure_count >= DUCKDUCKGO_BREAKER_FAILURE_THRESHOLD:
        _duckduckgo_open_until = time.monotonic() + DUCKDUCKGO_BREAKER_COOLDOWN_SECONDS


def _duckduckgo_should_log_breaker_open() -> bool:
    global _duckduckgo_last_open_log_at
    now = time.monotonic()
    if (
        now - _duckduckgo_last_open_log_at
        < DUCKDUCKGO_BREAKER_OPEN_LOG_INTERVAL_SECONDS
    ):
        return False
    _duckduckgo_last_open_log_at = now
    return True


def _duckduckgo_should_log_error() -> bool:
    global _duckduckgo_last_error_log_at
    now = time.monotonic()
    if now - _duckduckgo_last_error_log_at < DUCKDUCKGO_ERROR_LOG_INTERVAL_SECONDS:
        return False
    _duckduckgo_last_error_log_at = now
    return True


def _format_exception_summary(exc: BaseException) -> str:
    message = str(exc).strip()
    summary = type(exc).__name__
    if message:
        summary = f"{summary}: {message}"

    cause = exc.__cause__ or exc.__context__
    if cause is not None:
        cause_message = str(cause).strip()
        same_type = type(cause) is type(exc)
        same_message = cause_message == message
        if not (same_type and same_message):
            cause_summary = type(cause).__name__
            if cause_message:
                cause_summary = f"{cause_summary}: {cause_message}"
            summary = f"{summary} (cause={cause_summary})"

    return summary


REFERENCE_INLINE_PATTERN = re.compile(
    r"(?i)\b(references|bibliography|works\s+cited|tai\s+lieu\s+tham\s+khao)\s*:?\s*[\r\n]"
)

SOURCE_PRIORITY = {
    "crossref": 1.0,
    "pubmed": 0.95,
    "arxiv": 0.9,
    "core": 0.85,
    "vietnamese": 0.88,
    "duckduckgo": 0.75,
}

_query_candidate_cache: dict[str, list["NormalizedSourceCandidate"]] = {}

CITATION_STYLES = {
    "apa": re.compile(
        r"(?i)\b[A-Z][a-zA-ZÀ-ỹ]+(?:\s+[A-Z][a-zA-ZÀ-ỹ]+)*\s*\((\d{4})\)"
    ),
    "mla": re.compile(
        r"(?i)\b[A-Z][a-zA-ZÀ-ỹ]+(?:\s+[A-Z][a-zA-ZÀ-ỹ]+)*\s*,\s*.*?\((\d{4})\)"
    ),
    "ieee": re.compile(r"(?i)\[\s*(\d+)\s*\]"),
    "vietnamese": re.compile(
        r"(?i)\b[A-ZÀ-Ỹ][a-zà-ỹ]+(?:\s+[A-ZÀ-Ỹ][a-zà-ỹ]+)*\s*\(năm\s*(\d{4})\)"
    ),
}

VIETNAMESE_SOURCE_PATTERNS = [
    re.compile(r"scholar\.google\.com\.vn", re.IGNORECASE),
    re.compile(r"vjol\.info", re.IGNORECASE),
    re.compile(r"tailieu\.vn", re.IGNORECASE),
    re.compile(r"123doc\.org", re.IGNORECASE),
    re.compile(r"word\.com\.vn", re.IGNORECASE),
]

BIBLIOGRAPHY_PATTERNS = [
    re.compile(
        r"(?im)^\s*(tài\s+liệu\s+tham\s+khảo|bibliography|references|works\s+cited|literature\s+cited)\s*:?\s*$"
    ),
    re.compile(r"(?im)^\s*\[\d+\]\s+[A-Z][a-zA-ZÀ-ỹ].+"),
    re.compile(r"(?im)^\s*\(\d{4}\)\s+[A-Z][a-zA-ZÀ-ỹ].+"),
]

QUOTE_PATTERNS = [
    re.compile(r'"[^"]{10,500}"'),
    re.compile(r"'[^']{10,500}'"),
    re.compile(r"«[^»]{10,500}»"),
    re.compile(r"'''[^''']{10,500}'''"),
]

REFERENCE_PATTERNS = [
    re.compile(r"\[\d+\]|\(\w+,\s*\d{4}\)|\([a-zA-ZÀ-ỹ]+\s+et\s+al\.,\s*\d{4}\)"),
    re.compile(r"(?i)see\s+(also\s+)?(table|figure|fig\.)\s*\d+"),
    re.compile(r"(?i)as\s+cited\s+in"),
    re.compile(r"(?i)retrieved\s+from\s+https?://"),
]

EXCLUDE_SECTIONS = [
    re.compile(
        r"(?im)^(acknowledgements|acknowledgments|appendix|references|bibliography|tài\s+liệu\s+tham\s+khảo)\s*$"
    ),
    re.compile(r"(?im)^(figure|table|fig\.)\s+\d+"),
]


class CandidateIdentifiers(BaseModel):
    doi: str | None = None
    pmid: str | None = None
    arxiv_id: str | None = None


class NormalizedSourceCandidate(BaseModel):
    model_config = ConfigDict(extra="forbid")

    source: str = Field(..., min_length=1)
    canonical_url: str = Field(..., min_length=1)
    title: str = Field(..., min_length=1)
    snippet: str | None = None
    year: int | None = None
    authors: list[str] | None = None
    identifiers: CandidateIdentifiers = Field(default_factory=CandidateIdentifiers)


class SourceConnector(Protocol):
    name: str

    async def search(
        self,
        query: str,
        client: httpx.AsyncClient,
        limit: int,
    ) -> list[NormalizedSourceCandidate]: ...


class DuckDuckGoConnector:
    name = "duckduckgo"

    async def search(
        self,
        query: str,
        client: httpx.AsyncClient,
        limit: int,
    ) -> list[NormalizedSourceCandidate]:
        candidates: list[NormalizedSourceCandidate] = []
        search_query = quote(query[:200])
        ddg_url = f"https://html.duckduckgo.com/html/?q={search_query}"

        timeout = httpx.Timeout(
            timeout=settings.PLAGIARISM_SOURCE_TIMEOUT_SECONDS,
            connect=DUCKDUCKGO_CONNECT_TIMEOUT_SECONDS,
        )

        response: httpx.Response | None = None
        last_error: Exception | None = None
        for attempt in range(DUCKDUCKGO_CONNECT_MAX_RETRIES):
            try:
                response = await client.get(
                    ddg_url,
                    headers={"User-Agent": get_random_user_agent()},
                    timeout=timeout,
                )
                if (
                    response.status_code in DUCKDUCKGO_RETRYABLE_STATUS_CODES
                    and attempt < DUCKDUCKGO_CONNECT_MAX_RETRIES - 1
                ):
                    backoff = DUCKDUCKGO_BACKOFF_BASE_SECONDS * (2**attempt)
                    backoff += random.uniform(
                        0.0, DUCKDUCKGO_BACKOFF_JITTER_MAX_SECONDS
                    )
                    await asyncio.sleep(backoff)
                    continue
                break
            except (
                httpx.ConnectTimeout,
                httpx.ConnectError,
                httpx.TimeoutException,
            ) as e:
                last_error = e
                if attempt >= DUCKDUCKGO_CONNECT_MAX_RETRIES - 1:
                    raise
                backoff = DUCKDUCKGO_BACKOFF_BASE_SECONDS * (2**attempt)
                backoff += random.uniform(0.0, DUCKDUCKGO_BACKOFF_JITTER_MAX_SECONDS)
                await asyncio.sleep(backoff)

        if response is None:
            if last_error is not None:
                raise last_error
            return candidates

        if response.status_code != 200:
            return candidates

        matches = re.findall(r'uddg=([^"&]+)', response.text)
        from urllib.parse import unquote

        for match in matches:
            if len(candidates) >= limit:
                break
            try:
                url = unquote(match)
                if not url.startswith("http") or "duckduckgo.com" in url:
                    continue
                candidates.append(
                    NormalizedSourceCandidate(
                        source=self.name,
                        canonical_url=url,
                        title=url,
                    )
                )
            except Exception:
                continue

        return candidates


class VietnameseConnector:
    name = "vietnamese"
    VIETNAMESE_DOMAINS = [
        "vjol.info",
        "scholar.google.com.vn",
        "tailieu.vn",
        "123doc.org",
        "word.com.vn",
        "vanhoahoc.com",
        "text.edu.vn",
        "hocmai.vn",
    ]

    async def search(
        self,
        query: str,
        client: httpx.AsyncClient,
        limit: int,
    ) -> list[NormalizedSourceCandidate]:
        candidates: list[NormalizedSourceCandidate] = []

        for domain in self.VIETNAMESE_DOMAINS:
            if len(candidates) >= limit:
                break

            search_query = quote(f"{query[:100]} site:{domain}")
            ddg_url = f"https://html.duckduckgo.com/html/?q={search_query}"

            try:
                response = await client.get(
                    ddg_url,
                    headers={"User-Agent": get_random_user_agent()},
                    timeout=settings.PLAGIARISM_SOURCE_TIMEOUT_SECONDS,
                )

                if response.status_code != 200:
                    continue

                matches = re.findall(r'uddg=([^"&]+)', response.text)
                from urllib.parse import unquote

                for match in matches:
                    if len(candidates) >= limit:
                        break
                    try:
                        url = unquote(match)
                        if not url.startswith("http") or "duckduckgo.com" in url:
                            continue
                        if domain not in url:
                            continue
                        candidates.append(
                            NormalizedSourceCandidate(
                                source=self.name,
                                canonical_url=url,
                                title=url,
                                snippet=f"Vietnamese academic source from {domain}",
                            )
                        )
                    except Exception:
                        continue
            except Exception:
                continue

        return candidates


class CrossRefConnector:
    name = "crossref"

    async def search(
        self,
        query: str,
        client: httpx.AsyncClient,
        limit: int,
    ) -> list[NormalizedSourceCandidate]:
        candidates: list[NormalizedSourceCandidate] = []
        search_query = quote(query[:150])
        crossref_url = (
            f"https://api.crossref.org/works?query={search_query}&rows={limit}"
        )

        response = await client.get(
            crossref_url,
            headers={"User-Agent": get_random_user_agent()},
            timeout=settings.PLAGIARISM_SOURCE_TIMEOUT_SECONDS,
        )

        if response.status_code != 200:
            return candidates

        items = response.json().get("message", {}).get("items", [])
        for item in items:
            url = item.get("URL")
            if not url:
                continue

            title_value = item.get("title", [])
            title = (
                title_value[0] if isinstance(title_value, list) and title_value else url
            )
            published_parts = (
                item.get("published-print", {}).get("date-parts")
                or item.get("published-online", {}).get("date-parts")
                or []
            )
            year = None
            if (
                published_parts
                and isinstance(published_parts[0], list)
                and published_parts[0]
            ):
                first = published_parts[0][0]
                year = first if isinstance(first, int) else None

            authors = []
            for author in item.get("author", []):
                given = author.get("given")
                family = author.get("family")
                name = " ".join(v for v in [given, family] if v)
                if name:
                    authors.append(name)

            candidates.append(
                NormalizedSourceCandidate(
                    source=self.name,
                    canonical_url=url,
                    title=title,
                    snippet=item.get("abstract"),
                    year=year,
                    authors=authors or None,
                    identifiers=CandidateIdentifiers(doi=item.get("DOI")),
                )
            )

            if len(candidates) >= limit:
                break

        return candidates


class ArxivConnector:
    name = "arxiv"

    def __init__(self) -> None:
        self._last_request_at: float | None = None

    async def _enforce_politeness_delay(self) -> None:
        if self._last_request_at is None:
            return

        now = asyncio.get_running_loop().time()
        elapsed = now - self._last_request_at
        remaining = ARXIV_POLITENESS_DELAY_SECONDS - elapsed
        if remaining > 0:
            await asyncio.sleep(remaining)

    async def _request_with_retry(
        self,
        client: httpx.AsyncClient,
        params: dict[str, str | int],
    ) -> httpx.Response | None:
        total_attempts = ARXIV_MAX_RETRIES + 1
        for attempt in range(total_attempts):
            await self._enforce_politeness_delay()
            self._last_request_at = asyncio.get_running_loop().time()

            try:
                response = await client.get(
                    ARXIV_API_URL,
                    params=params,
                    headers={"User-Agent": get_random_user_agent()},
                    timeout=ARXIV_TIMEOUT_SECONDS,
                )
            except Exception:
                response = None

            if response is not None and response.status_code == 200:
                return response

            if attempt < total_attempts - 1:
                backoff_base = ARXIV_BACKOFF_SECONDS[
                    min(attempt, len(ARXIV_BACKOFF_SECONDS) - 1)
                ]
                await asyncio.sleep(
                    backoff_base + random.uniform(0.0, ARXIV_BACKOFF_JITTER_MAX_SECONDS)
                )

        return None

    async def search(
        self,
        query: str,
        client: httpx.AsyncClient,
        limit: int,
    ) -> list[NormalizedSourceCandidate]:
        bounded_limit = max(0, min(int(limit), ARXIV_MAX_RESULTS_UPPER_BOUND))
        bounded_start = max(0, min(0, ARXIV_MAX_START_WINDOW))
        if bounded_limit <= 0:
            return []

        params: dict[str, str | int] = {
            "search_query": f"all:{query[:200]}",
            "start": bounded_start,
            "max_results": bounded_limit,
            "sortBy": "relevance",
            "sortOrder": "descending",
        }

        response = await self._request_with_retry(client=client, params=params)
        if response is None:
            return []

        try:
            root = ET.fromstring(response.text)
        except ET.ParseError:
            return []

        atom_ns = "{http://www.w3.org/2005/Atom}"
        entries = root.findall(f"{atom_ns}entry")

        candidates: list[NormalizedSourceCandidate] = []
        for entry in entries:
            if len(candidates) >= bounded_limit:
                break

            id_text = entry.findtext(f"{atom_ns}id")
            title_text = entry.findtext(f"{atom_ns}title")
            if id_text is None or title_text is None:
                continue

            raw_url = id_text.strip()
            if not raw_url:
                continue

            canonical_url = raw_url.replace("http://", "https://")
            if "/pdf/" in canonical_url:
                canonical_url = canonical_url.replace("/pdf/", "/abs/")

            if "/abs/" not in canonical_url:
                continue

            arxiv_id = canonical_url.rstrip("/").split("/abs/")[-1]
            if not arxiv_id:
                continue

            summary_text = entry.findtext(f"{atom_ns}summary")
            snippet = summary_text.strip() if summary_text else None

            year = None
            published_text = entry.findtext(f"{atom_ns}published")
            if published_text:
                year_match = re.match(r"^(\d{4})", published_text.strip())
                if year_match:
                    year = int(year_match.group(1))

            authors = []
            for author in entry.findall(f"{atom_ns}author"):
                name_text = author.findtext(f"{atom_ns}name")
                if name_text and name_text.strip():
                    authors.append(name_text.strip())

            title = title_text.strip()
            if not title:
                continue

            candidates.append(
                NormalizedSourceCandidate(
                    source=self.name,
                    canonical_url=canonical_url,
                    title=title,
                    snippet=snippet,
                    year=year,
                    authors=authors or None,
                    identifiers=CandidateIdentifiers(arxiv_id=arxiv_id),
                )
            )

        return candidates


class CoreConnector:
    name = "core"

    def __init__(self) -> None:
        self._last_request_at: float | None = None
        self._request_semaphore = asyncio.Semaphore(CORE_CONNECTOR_CONCURRENCY)
        self._consecutive_failures = 0
        self._degraded_until: float | None = None

    def _api_key(self) -> str:
        return settings.CORE_API_KEY.strip()

    def _auth_mode(self) -> str:
        mode = settings.PLAGIARISM_SOURCE_CORE_AUTH_MODE.strip().lower()
        if mode == CORE_AUTH_MODE_QUERY:
            return CORE_AUTH_MODE_QUERY
        return CORE_AUTH_MODE_BEARER

    def _build_auth(self) -> tuple[dict[str, str], dict[str, str]]:
        api_key = self._api_key()
        if not api_key:
            return {}, {}

        if self._auth_mode() == CORE_AUTH_MODE_QUERY:
            return {}, {"api_key": api_key}

        return {"Authorization": f"Bearer {api_key}"}, {}

    async def _enforce_throttle(self) -> None:
        if self._last_request_at is None:
            return

        now = asyncio.get_running_loop().time()
        elapsed = now - self._last_request_at
        remaining = CORE_MIN_REQUEST_INTERVAL_SECONDS - elapsed
        if remaining > 0:
            await asyncio.sleep(remaining)

    def _is_degraded(self) -> bool:
        if self._degraded_until is None:
            return False
        return asyncio.get_running_loop().time() < self._degraded_until

    def _record_success(self) -> None:
        self._consecutive_failures = 0
        self._degraded_until = None

    def _record_failure(self) -> None:
        self._consecutive_failures += 1
        if self._consecutive_failures >= CORE_DEGRADED_FAILURE_THRESHOLD:
            now = asyncio.get_running_loop().time()
            self._degraded_until = now + CORE_DEGRADED_WINDOW_SECONDS
            self._consecutive_failures = 0

    def _parse_retry_after_seconds(self, response: httpx.Response) -> float | None:
        retry_after = response.headers.get("Retry-After")
        if not retry_after:
            return None

        try:
            return max(0.0, float(retry_after))
        except ValueError:
            pass

        try:
            retry_time = parsedate_to_datetime(retry_after)
        except (TypeError, ValueError):
            return None

        if retry_time.tzinfo is None:
            retry_time = retry_time.replace(tzinfo=timezone.utc)

        current_utc = datetime.now(timezone.utc)
        return max(0.0, (retry_time - current_utc).total_seconds())

    def _should_retry_response(self, response: httpx.Response) -> bool:
        return response.status_code == 429 or response.status_code >= 500

    def _year_from_published(self, published_date: str | None) -> int | None:
        if not published_date:
            return None
        match = re.match(r"^(\d{4})", str(published_date).strip())
        return int(match.group(1)) if match else None

    def _authors_from_item(self, item: dict) -> list[str] | None:
        authors = []
        for author in item.get("authors", []):
            if isinstance(author, dict):
                name = str(author.get("name") or "").strip()
            else:
                name = str(author).strip()
            if name:
                authors.append(name)
        return authors or None

    def _canonical_url_from_item(self, item: dict, doi: str | None) -> str | None:
        for key in ("downloadUrl", "sourceFulltextUrls", "links"):
            value = item.get(key)
            if isinstance(value, str) and value.startswith("http"):
                return value
            if isinstance(value, list):
                for nested in value:
                    if isinstance(nested, str) and nested.startswith("http"):
                        return nested
                    if isinstance(nested, dict):
                        nested_url = nested.get("url")
                        if isinstance(nested_url, str) and nested_url.startswith(
                            "http"
                        ):
                            return nested_url
        if doi:
            return f"https://doi.org/{doi}"
        core_id = item.get("id")
        if core_id is not None:
            return f"https://core.ac.uk/works/{core_id}"
        return None

    def _normalize_item(self, item: dict) -> NormalizedSourceCandidate | None:
        title = str(item.get("title") or "").strip()
        if not title:
            return None

        doi = str(item.get("doi") or "").strip() or None
        canonical_url = self._canonical_url_from_item(item, doi)
        if not canonical_url:
            return None

        snippet = item.get("abstract") or item.get("description")
        snippet_text = str(snippet).strip() if isinstance(snippet, str) else None

        return NormalizedSourceCandidate(
            source=self.name,
            canonical_url=canonical_url,
            title=title,
            snippet=snippet_text,
            year=self._year_from_published(item.get("publishedDate")),
            authors=self._authors_from_item(item),
            identifiers=CandidateIdentifiers(doi=doi),
        )

    async def _request_with_retry(
        self,
        client: httpx.AsyncClient,
        params: dict[str, str | int],
        headers: dict[str, str],
    ) -> httpx.Response | None:
        total_attempts = CORE_MAX_RETRIES + 1
        for attempt in range(total_attempts):
            await self._enforce_throttle()
            self._last_request_at = asyncio.get_running_loop().time()

            try:
                response = await client.get(
                    CORE_API_URL,
                    params=params,
                    headers=headers,
                    timeout=CORE_TIMEOUT_SECONDS,
                )
            except (httpx.TimeoutException, httpx.NetworkError, httpx.RequestError):
                response = None

            if response is not None:
                if response.status_code == 200:
                    return response
                if not self._should_retry_response(response):
                    return None

            if attempt >= total_attempts - 1:
                break

            retry_after_seconds = (
                self._parse_retry_after_seconds(response)
                if response is not None
                else None
            )
            backoff_base = CORE_BACKOFF_SECONDS[
                min(attempt, len(CORE_BACKOFF_SECONDS) - 1)
            ]
            fallback_backoff = backoff_base + random.uniform(
                0.0, CORE_BACKOFF_JITTER_MAX_SECONDS
            )
            await asyncio.sleep(
                retry_after_seconds
                if retry_after_seconds is not None
                else fallback_backoff
            )

        return None

    async def search(
        self,
        query: str,
        client: httpx.AsyncClient,
        limit: int,
    ) -> list[NormalizedSourceCandidate]:
        bounded_limit = max(0, int(limit))
        if bounded_limit <= 0:
            return []

        if self._is_degraded():
            return []

        if not self._api_key():
            return []

        headers, auth_query = self._build_auth()
        params: dict[str, str | int] = {
            "q": query[:200],
            "limit": bounded_limit,
            **auth_query,
        }

        async with self._request_semaphore:
            response = await self._request_with_retry(client, params, headers)

        if response is None:
            self._record_failure()
            return []

        self._record_success()

        results = response.json().get("results", [])
        candidates: list[NormalizedSourceCandidate] = []
        for item in results:
            if len(candidates) >= bounded_limit:
                break
            if not isinstance(item, dict):
                continue
            normalized = self._normalize_item(item)
            if normalized is None:
                continue
            candidates.append(normalized)

        return candidates


class PubMedConnector:
    name = "pubmed"

    def __init__(self) -> None:
        self._last_request_at: float | None = None
        self._request_semaphore = asyncio.Semaphore(PUBMED_CONNECTOR_CONCURRENCY)

    def _api_key(self) -> str:
        return settings.NCBI_API_KEY.strip()

    def _identity_params(self) -> dict[str, str]:
        params = {
            "tool": settings.NCBI_TOOL,
            "email": settings.NCBI_EMAIL,
        }
        api_key = self._api_key()
        if api_key:
            params["api_key"] = api_key
        return params

    def _throttle_interval(self) -> float:
        return (
            PUBMED_KEY_MIN_INTERVAL_SECONDS
            if self._api_key()
            else PUBMED_NO_KEY_MIN_INTERVAL_SECONDS
        )

    async def _enforce_throttle(self) -> None:
        if self._last_request_at is None:
            return

        now = asyncio.get_running_loop().time()
        elapsed = now - self._last_request_at
        remaining = self._throttle_interval() - elapsed
        if remaining > 0:
            await asyncio.sleep(remaining)

    def _parse_retry_after_seconds(self, response: httpx.Response) -> float | None:
        retry_after = response.headers.get("Retry-After")
        if not retry_after:
            return None

        try:
            return max(0.0, float(retry_after))
        except ValueError:
            pass

        try:
            retry_time = parsedate_to_datetime(retry_after)
        except (TypeError, ValueError):
            return None

        if retry_time.tzinfo is None:
            retry_time = retry_time.replace(tzinfo=timezone.utc)

        current_utc = datetime.now(timezone.utc)
        return max(0.0, (retry_time - current_utc).total_seconds())

    def _should_retry_response(self, response: httpx.Response) -> bool:
        return response.status_code == 429 or response.status_code >= 500

    async def _request_with_retry(
        self,
        client: httpx.AsyncClient,
        url: str,
        params: dict[str, str | int],
    ) -> httpx.Response | None:
        total_attempts = PUBMED_MAX_RETRIES + 1
        for attempt in range(total_attempts):
            await self._request_semaphore.acquire()
            try:
                await self._enforce_throttle()
                self._last_request_at = asyncio.get_running_loop().time()
                try:
                    response = await client.get(
                        url,
                        params=params,
                        timeout=settings.PLAGIARISM_SOURCE_TIMEOUT_SECONDS,
                    )
                except (httpx.TimeoutException, httpx.NetworkError, httpx.RequestError):
                    response = None
            finally:
                self._request_semaphore.release()

            if response is not None:
                if response.status_code == 200:
                    return response
                if not self._should_retry_response(response):
                    return None

            if attempt >= total_attempts - 1:
                break

            retry_after_seconds = (
                self._parse_retry_after_seconds(response)
                if response is not None
                else None
            )
            backoff_base = PUBMED_BACKOFF_SECONDS[
                min(attempt, len(PUBMED_BACKOFF_SECONDS) - 1)
            ]
            fallback_backoff = backoff_base + random.uniform(
                0.0, PUBMED_BACKOFF_JITTER_MAX_SECONDS
            )
            await asyncio.sleep(
                retry_after_seconds
                if retry_after_seconds is not None
                else fallback_backoff
            )

        return None

    async def _search_uids_page(
        self,
        query: str,
        client: httpx.AsyncClient,
        *,
        retmax: int,
        retstart: int,
    ) -> list[str]:
        params: dict[str, str | int] = {
            "db": "pubmed",
            "term": query[:200],
            "retmode": "json",
            "sort": "relevance",
            "retmax": retmax,
            "retstart": retstart,
            **self._identity_params(),
        }
        response = await self._request_with_retry(client, PUBMED_ESEARCH_URL, params)
        if response is None:
            return []

        id_list = response.json().get("esearchresult", {}).get("idlist", [])
        return [str(uid) for uid in id_list if str(uid).isdigit()]

    async def _fetch_summaries(
        self,
        client: httpx.AsyncClient,
        pmids: list[str],
    ) -> dict[str, dict]:
        if not pmids:
            return {}

        params: dict[str, str | int] = {
            "db": "pubmed",
            "id": ",".join(pmids),
            "retmode": "json",
            **self._identity_params(),
        }
        response = await self._request_with_retry(client, PUBMED_ESUMMARY_URL, params)
        if response is None:
            return {}

        payload = response.json().get("result", {})
        return {
            pmid: payload.get(pmid, {})
            for pmid in payload.get("uids", [])
            if isinstance(payload.get(pmid), dict)
        }

    def _year_from_pubdate(self, pubdate: str | None) -> int | None:
        if not pubdate:
            return None
        match = re.search(r"(19|20)\d{2}", pubdate)
        return int(match.group(0)) if match else None

    def _metadata_snippet(self, summary: dict) -> str | None:
        source_name = summary.get("source")
        pubdate = summary.get("pubdate")
        if source_name and pubdate:
            return f"{source_name} ({pubdate})"
        if source_name:
            return str(source_name)
        return None

    def _normalize_summary(
        self, pmid: str, summary: dict
    ) -> NormalizedSourceCandidate | None:
        title = str(summary.get("title") or "").strip()
        if not title:
            return None

        authors = []
        for author in summary.get("authors", []):
            name = str(author.get("name") or "").strip()
            if name:
                authors.append(name)

        return NormalizedSourceCandidate(
            source=self.name,
            canonical_url=f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/",
            title=title,
            snippet=self._metadata_snippet(summary),
            year=self._year_from_pubdate(summary.get("pubdate")),
            authors=authors or None,
            identifiers=CandidateIdentifiers(pmid=pmid),
        )

    async def search(
        self,
        query: str,
        client: httpx.AsyncClient,
        limit: int,
    ) -> list[NormalizedSourceCandidate]:
        bounded_limit = max(0, min(int(limit), PUBMED_RETMAX_UPPER_BOUND))
        if bounded_limit <= 0:
            return []

        page_size = min(10, bounded_limit)
        max_retstart = min(PUBMED_MAX_START_WINDOW, PUBMED_RETMAX_UPPER_BOUND)
        retstart = 0
        seen_pmids: set[str] = set()
        ordered_pmids: list[str] = []

        while len(ordered_pmids) < bounded_limit and retstart <= max_retstart:
            page_uids = await self._search_uids_page(
                query,
                client,
                retmax=page_size,
                retstart=retstart,
            )
            if not page_uids:
                break

            for uid in page_uids:
                if uid in seen_pmids:
                    continue
                seen_pmids.add(uid)
                ordered_pmids.append(uid)
                if len(ordered_pmids) >= bounded_limit:
                    break

            if len(page_uids) < page_size:
                break

            retstart += page_size

        summaries = await self._fetch_summaries(client, ordered_pmids[:bounded_limit])
        candidates: list[NormalizedSourceCandidate] = []
        for pmid in ordered_pmids:
            if len(candidates) >= bounded_limit:
                break
            summary = summaries.get(pmid)
            if not summary:
                continue
            candidate = self._normalize_summary(pmid, summary)
            if candidate is None:
                continue
            candidates.append(candidate)

        return candidates


DUCKDUCKGO_CONNECTOR = DuckDuckGoConnector()
CROSSREF_CONNECTOR = CrossRefConnector()
ARXIV_CONNECTOR = ArxivConnector()
CORE_CONNECTOR = CoreConnector()
PUBMED_CONNECTOR = PubMedConnector()
VIETNAMESE_CONNECTOR = VietnameseConnector()


class OpenAlexConnector:
    name = "openalex"

    async def search(
        self,
        query: str,
        client: httpx.AsyncClient,
        limit: int,
    ) -> list[NormalizedSourceCandidate]:
        candidates: list[NormalizedSourceCandidate] = []
        search_query = quote(query[:150])

        openalex_url = (
            f"https://api.openalex.org/works?search={search_query}&per_page={limit}"
        )

        headers = {"User-Agent": get_random_user_agent()}
        if settings.OPENALEX_EMAIL:
            headers["mailto"] = settings.OPENALEX_EMAIL

        try:
            response = await client.get(
                openalex_url,
                headers=headers,
                timeout=settings.PLAGIARISM_SOURCE_TIMEOUT_SECONDS,
            )

            if response.status_code != 200:
                return candidates

            data = response.json()
            items = data.get("results", [])

            for item in items:
                doi = item.get("doi")
                url = (
                    item.get("display_url")
                    or item.get("doi")
                    or f"https://openalex.org/{item.get('id', '').split('/')[-1]}"
                )

                title = item.get("title") or "Untitled"

                year = None
                if item.get("publication_year"):
                    year = item.get("publication_year")

                authors = []
                for author in item.get("authorships", [])[:5]:
                    author_name = author.get("author", {}).get("display_name")
                    if author_name:
                        authors.append(author_name)

                abstract = item.get("abstract_inverted_index")
                abstract_text = None
                if abstract:
                    words = []
                    for word, positions in sorted(
                        abstract.items(),
                        key=lambda x: min(x[1]) if x[1] else float("inf"),
                    ):
                        words.append(word)
                    abstract_text = " ".join(words[:200])

                candidates.append(
                    NormalizedSourceCandidate(
                        source=self.name,
                        canonical_url=url,
                        title=title,
                        snippet=abstract_text,
                        year=year,
                        authors=authors or None,
                        identifiers=CandidateIdentifiers(doi=doi),
                    )
                )

                if len(candidates) >= limit:
                    break

        except Exception as e:
            print(f"OpenAlex search error: {e}")

        return candidates


OPENALEX_CONNECTOR = OpenAlexConnector()


def normalize_candidate_payload(
    payload: dict | NormalizedSourceCandidate,
) -> NormalizedSourceCandidate:
    if isinstance(payload, NormalizedSourceCandidate):
        return payload
    return NormalizedSourceCandidate.model_validate(payload)


def get_source_connector_registry() -> list[SourceConnector]:
    connectors: list[SourceConnector] = []
    if settings.PLAGIARISM_SOURCE_DUCKDUCKGO_ENABLED:
        connectors.append(DUCKDUCKGO_CONNECTOR)
    if settings.PLAGIARISM_SOURCE_CROSSREF_ENABLED:
        connectors.append(CROSSREF_CONNECTOR)
    if settings.PLAGIARISM_SOURCE_ARXIV_ENABLED:
        connectors.append(ARXIV_CONNECTOR)
    if settings.PLAGIARISM_SOURCE_CORE_ENABLED:
        connectors.append(CORE_CONNECTOR)
    if settings.PLAGIARISM_SOURCE_PUBMED_ENABLED:
        connectors.append(PUBMED_CONNECTOR)
    if settings.PLAGIARISM_SOURCE_VIETNAMESE_ENABLED:
        connectors.append(VIETNAMESE_CONNECTOR)
    if settings.PLAGIARISM_SOURCE_OPENALEX_ENABLED:
        connectors.append(OPENALEX_CONNECTOR)
    return connectors


def get_source_caps() -> dict[str, int]:
    return {
        "duckduckgo": max(0, int(settings.PLAGIARISM_SOURCE_DUCKDUCKGO_MAX_CANDIDATES)),
        "crossref": max(0, int(settings.PLAGIARISM_SOURCE_CROSSREF_MAX_CANDIDATES)),
        "arxiv": max(0, int(settings.PLAGIARISM_SOURCE_ARXIV_MAX_CANDIDATES)),
        "core": max(0, int(settings.PLAGIARISM_SOURCE_CORE_MAX_CANDIDATES)),
        "pubmed": max(0, int(settings.PLAGIARISM_SOURCE_PUBMED_MAX_CANDIDATES)),
        "vietnamese": max(0, int(settings.PLAGIARISM_SOURCE_VIETNAMESE_MAX_CANDIDATES)),
        "openalex": max(0, int(settings.PLAGIARISM_SOURCE_OPENALEX_MAX_CANDIDATES)),
    }


def _normalize_doi(doi: str | None) -> str | None:
    if not doi:
        return None
    normalized = doi.strip().lower()
    normalized = re.sub(r"^https?://(dx\.)?doi\.org/", "", normalized)
    return normalized or None


def _extract_doi_from_url(url: str) -> str | None:
    lowered = url.strip().lower()
    if "doi.org/" not in lowered:
        return None
    return _normalize_doi(lowered.split("doi.org/", maxsplit=1)[-1])


def _normalize_canonical_url(url: str) -> str:
    """Normalize URL: remove tracking params, trailing slash, force https."""
    try:
        from urllib.parse import urlparse, parse_qs, urlencode, urlunparse
    except ImportError:
        url = url.strip().lower()
        url = url.replace("http://", "https://", 1)
        return url.rstrip("/")

    # Tracking parameters to remove
    TRACKING_PARAMS = [
        # Google
        "utm_source",
        "utm_medium",
        "utm_campaign",
        "utm_term",
        "utm_content",
        "gclid",
        "gclsrc",
        "dclid",
        # Facebook
        "fbclid",
        "fb_action_ids",
        "fb_action_types",
        "fb_source",
        # Twitter
        "twclid",
        # Generic
        "mc_cid",
        "mc_eid",  # Mailchimp
        "ref",
        "ref_src",
        "ref_url",
        "source",
        "affiliate",
        "_ga",
        "_gl",
        "si",
        "igshid",  # Various
    ]

    try:
        parsed = urlparse(url.strip().lower())

        # Remove tracking params
        query_params = parse_qs(parsed.query, keep_blank_values=True)
        clean_params = {
            k: v for k, v in query_params.items() if k.lower() not in TRACKING_PARAMS
        }

        # Rebuild URL
        clean_query = urlencode(clean_params, doseq=True)
        normalized = urlunparse(
            (
                parsed.scheme or "https",
                parsed.netloc,
                parsed.path,
                parsed.params,
                clean_query,
                "",
            )
        )

        return normalized.rstrip("/")
    except Exception:
        # Fallback
        url = url.strip().lower()
        url = url.replace("http://", "https://", 1)
        return url.rstrip("/")


def _normalize_title(title: str) -> str:
    collapsed = re.sub(r"\s+", " ", title.strip().lower())
    return re.sub(r"[^a-z0-9\s]", "", collapsed)


def _dedupe_keys(candidate: NormalizedSourceCandidate) -> list[tuple[str, str]]:
    keys: list[tuple[str, str]] = []

    doi = _normalize_doi(candidate.identifiers.doi) or _extract_doi_from_url(
        candidate.canonical_url
    )
    if doi:
        keys.append(("doi", doi))

    pmid = (candidate.identifiers.pmid or "").strip()
    if pmid:
        keys.append(("pmid", pmid))

    arxiv_id = (candidate.identifiers.arxiv_id or "").strip().lower()
    if arxiv_id:
        keys.append(("arxiv", arxiv_id))

    canonical = _normalize_canonical_url(candidate.canonical_url)
    if canonical:
        keys.append(("url", canonical))

    title = _normalize_title(candidate.title)
    if title:
        keys.append(("title", title))

    return keys


def _duplicate_quality_score(candidate: NormalizedSourceCandidate) -> float:
    score = SOURCE_PRIORITY.get(candidate.source.lower(), 0.5)
    if candidate.snippet and candidate.snippet.strip():
        score += 0.8
    if candidate.year is not None:
        score += 0.4
    if candidate.authors:
        score += min(0.6, len(candidate.authors) * 0.15)
    if _normalize_doi(candidate.identifiers.doi):
        score += 1.5
    if candidate.identifiers.pmid:
        score += 1.1
    if candidate.identifiers.arxiv_id:
        score += 0.9
    return score


def _rank_candidate(query: str, candidate: NormalizedSourceCandidate) -> float:
    query_text = query.strip()
    blended = f"{candidate.title} {candidate.snippet or ''}".strip()
    lexical = max(
        calculate_cosine_similarity(query_text, blended),
        calculate_ngram_similarity(query_text, blended, n=3),
    )
    recency = 0.0
    if candidate.year is not None:
        recency = min(0.15, max(0.0, (candidate.year - 1990) / 300.0))
    metadata = 0.15 if candidate.snippet else 0.0
    source_weight = SOURCE_PRIORITY.get(candidate.source.lower(), 0.5) * 0.2
    return lexical + recency + metadata + source_weight


def _cache_query_candidates(
    query: str, candidates: list[NormalizedSourceCandidate]
) -> None:
    cache_key = query.strip()
    if not cache_key:
        return
    _query_candidate_cache[cache_key] = candidates
    if len(_query_candidate_cache) > SEARCH_RESULT_CACHE_MAX_ENTRIES:
        oldest_key = next(iter(_query_candidate_cache))
        _query_candidate_cache.pop(oldest_key, None)


def _get_cached_query_candidates(query: str) -> list[NormalizedSourceCandidate] | None:
    cache_key = query.strip()
    if not cache_key:
        return None
    cached = _query_candidate_cache.get(cache_key)
    return list(cached) if cached is not None else None


def _source_from_url(url: str) -> str:
    lowered = url.strip().lower()
    if "arxiv.org" in lowered:
        return "arxiv"
    if "pubmed.ncbi.nlm.nih.gov" in lowered:
        return "pubmed"
    if "doi.org" in lowered:
        return "doi"
    if "core.ac.uk" in lowered:
        return "core"
    return "web"


def _build_source_counts(results: list[SentenceResult]) -> dict[str, int] | None:
    counts: dict[str, int] = {}
    for sentence_result in results:
        for source in sentence_result.sources:
            source_name = _source_from_url(source.url)
            counts[source_name] = counts.get(source_name, 0) + 1
    return counts or None


def _source_group_key(url: str) -> str:
    normalized_url = _normalize_canonical_url(url)
    doi = _extract_doi_from_url(normalized_url)
    if doi:
        return f"doi:{doi}"
    return f"url:{normalized_url}"


def _build_sentence_offsets(results: list[SentenceResult]) -> list[tuple[int, int]]:
    offsets: list[tuple[int, int]] = []
    cursor = 0
    for sentence_result in results:
        start_char = cursor
        end_char = start_char + len(sentence_result.sentence)
        offsets.append((start_char, end_char))
        cursor = end_char + 1
    return offsets


def _build_report_v2_source_groups(
    results: list[SentenceResult],
) -> list[ReportV2SourceGroup]:
    sentence_offsets = _build_sentence_offsets(results)

    grouped: dict[
        str,
        dict[
            str,
            str | dict[int, ReportV2SourceSpan],
        ],
    ] = {}

    for sentence_index, sentence_result in enumerate(results):
        start_char, end_char = sentence_offsets[sentence_index]
        for source in sentence_result.sources:
            source_key = _source_group_key(source.url)
            source_entry = grouped.setdefault(
                source_key,
                {
                    "canonical_url": _normalize_canonical_url(source.url),
                    "source_type": _source_from_url(source.url),
                    "spans": {},
                },
            )

            sentence_spans = source_entry["spans"]
            assert isinstance(sentence_spans, dict)

            current_span = sentence_spans.get(sentence_index)
            candidate_span = ReportV2SourceSpan(
                sentence_index=sentence_index,
                start_char=start_char,
                end_char=end_char,
                similarity=source.similarity,
            )
            if (
                current_span is None
                or candidate_span.similarity > current_span.similarity
            ):
                sentence_spans[sentence_index] = candidate_span

    source_groups: list[ReportV2SourceGroup] = []
    for index, source_key in enumerate(sorted(grouped.keys()), start=1):
        entry = grouped[source_key]
        sentence_spans = entry["spans"]
        assert isinstance(sentence_spans, dict)
        spans = [sentence_spans[idx] for idx in sorted(sentence_spans.keys())]

        source_groups.append(
            ReportV2SourceGroup(
                source_id=f"src-{index:03d}",
                source_type=str(entry["source_type"]),
                canonical_url=str(entry["canonical_url"]),
                spans=spans,
            )
        )

    return source_groups


def _derive_confidence_band(
    *,
    total_sentences: int,
    sentences_with_sources: int,
    semantic_sentences: int,
    fallback_sentences: int,
    total_source_matches: int,
) -> str:
    if total_sentences <= 0:
        return "low"

    source_evidence_ratio = sentences_with_sources / total_sentences
    semantic_coverage_ratio = semantic_sentences / total_sentences
    fallback_ratio = fallback_sentences / total_sentences
    average_sources = total_source_matches / total_sentences

    score = 0
    if source_evidence_ratio >= 0.7:
        score += 2
    elif source_evidence_ratio >= 0.4:
        score += 1

    if semantic_coverage_ratio >= 0.5:
        score += 2
    elif semantic_coverage_ratio > 0.0:
        score += 1

    if fallback_ratio > 0.5:
        score -= 2
    elif fallback_ratio > 0.0:
        score -= 1

    if average_sources >= 1.5:
        score += 1

    if score >= 3:
        return "high"
    if score >= 1:
        return "medium"
    return "low"


CITATION_PATTERNS = [
    re.compile(r"\([A-Z][a-zA-Z]+(?:\s+et\s+al\.?)?(?:\s*,\s*\d{4}|\s+\d{4})?\)"),
    re.compile(r"\[\d+(?:,\s*\d+)*\]"),
    re.compile(r"¹|²|³|⁴|⁵|⁶|⁷|⁸|⁹|⁰"),
    re.compile(r"\([A-Z][a-zA-Z]+\s+and\s+others?,\s*\d{4}\)"),
    re.compile(r"\(\d{4}\)"),
]

QUOTATION_PATTERNS = [
    re.compile(r'"[^"]+"'),
    re.compile(r"«[^»]+»"),
    re.compile(r"『[^』]+』"),
    re.compile(r"「[^」]+」"),
    re.compile(r"''[^'']+''"),
]


def _has_citation(text: str) -> bool:
    for pattern in CITATION_PATTERNS:
        if pattern.search(text):
            return True
    return False


def _has_quotation(text: str) -> bool:
    for pattern in QUOTATION_PATTERNS:
        if pattern.search(text):
            return True
    return False


def _build_match_groups(results: list[SentenceResult]) -> list[MatchGroup]:
    groups: dict[str, list[str]] = {
        "not_cited_or_quoted": [],
        "missing_quotations": [],
        "missing_citation": [],
        "cited_and_quoted": [],
    }

    for sentence_result in results:
        if not sentence_result.sources:
            continue

        sentence = sentence_result.sentence
        has_citation = _has_citation(sentence)
        has_quotation = _has_quotation(sentence)

        if has_citation and has_quotation:
            groups["cited_and_quoted"].append(sentence)
        elif has_citation and not has_quotation:
            groups["missing_quotations"].append(sentence)
        elif not has_citation and has_quotation:
            groups["missing_citation"].append(sentence)
        else:
            groups["not_cited_or_quoted"].append(sentence)

    total = sum(len(matches) for matches in groups.values())
    if total == 0:
        return []

    match_groups = []
    for group_type, sentences in groups.items():
        count = len(sentences)
        percentage = round((count / total) * 100, 1) if total > 0 else 0.0
        sample_sentences = sentences[:3]

        match_groups.append(
            MatchGroup(
                group_type=group_type,
                count=count,
                percentage=percentage,
                sample_sentences=sample_sentences,
            )
        )

    return match_groups


def _filter_small_matches(
    results: list[SentenceResult],
    min_word_count: int,
) -> tuple[list[SentenceResult], int]:
    if min_word_count <= 0:
        return results, 0

    filtered = []
    removed_count = 0

    for result in results:
        word_count = len(result.sentence.split())
        if word_count >= min_word_count:
            filtered.append(result)
        else:
            removed_count += 1

    return filtered, removed_count


def _filter_small_sources(
    results: list[SentenceResult],
    min_source_count: int,
) -> list[SentenceResult]:
    if not min_source_count:
        return results

    filtered = []
    for result in results:
        if len(result.sources) >= min_source_count:
            filtered.append(result)

    return filtered


def _filter_by_source_type(
    results: list[SentenceResult],
    allowed_types: list[str],
) -> list[SentenceResult]:
    """Filter results by source type (web, academic, preprint)."""
    if not allowed_types:
        return results

    allowed_lower = [t.lower() for t in allowed_types]

    filtered = []
    for result in results:
        filtered_sources = []
        for source in result.sources:
            url = source.url.lower()
            source_type = "web"
            if any(x in url for x in ["arxiv", "preprint", "doi.org/10."]):
                source_type = "preprint"
            elif any(
                x in url
                for x in [
                    "scholar.google",
                    "pubmed",
                    "crossref",
                    "core.edu",
                    "semantic",
                ]
            ):
                source_type = "academic"

            if source_type in allowed_lower:
                filtered_sources.append(source)

        if filtered_sources:
            result.sources = filtered_sources
            filtered.append(result)

    return filtered


def _filter_by_contribution(
    results: list[SentenceResult],
    min_contribution_percent: int,
) -> list[SentenceResult]:
    """Filter out sources contributing less than N% to similarity."""
    if min_contribution_percent <= 0:
        return results

    filtered = []
    for result in results:
        if not result.sources:
            continue

        total_similarity = sum(s.similarity for s in result.sources)
        if total_similarity == 0:
            continue

        filtered_sources = []
        for source in result.sources:
            contribution = (source.similarity / total_similarity) * 100
            if contribution >= min_contribution_percent:
                filtered_sources.append(source)

        if filtered_sources:
            result.sources = filtered_sources
            filtered.append(result)

    return filtered


def _build_report_v2(
    results: list[SentenceResult],
    metadata_overrides: dict[str, str] | None = None,
    extra_caveats: list[ReportV2Caveat] | None = None,
) -> ReportV2:
    total_sentences = len(results)
    sentences_with_sources = sum(1 for r in results if r.sources)
    semantic_sentences = sum(1 for r in results if r.used_ai)
    fallback_sentences = sum(1 for r in results if r.fallback_used)
    total_source_matches = sum(len(r.sources) for r in results)
    source_groups = _build_report_v2_source_groups(results)
    match_groups = _build_match_groups(results)
    total_group_spans = sum(len(group.spans) for group in source_groups)

    confidence_band = _derive_confidence_band(
        total_sentences=total_sentences,
        sentences_with_sources=sentences_with_sources,
        semantic_sentences=semantic_sentences,
        fallback_sentences=fallback_sentences,
        total_source_matches=total_source_matches,
    )

    caveats: list[ReportV2Caveat] = []
    if fallback_sentences > 0:
        caveats.append(
            ReportV2Caveat(
                code="SEMANTIC_UNAVAILABLE_FALLBACK",
                message=(
                    "Semantic similarity was unavailable or exhausted for part of this "
                    "request; keyword fallback scoring was applied."
                ),
            )
        )

    if total_sentences > 0 and sentences_with_sources == 0:
        caveats.append(
            ReportV2Caveat(
                code="LIMITED_SOURCE_EVIDENCE",
                message=(
                    "No corroborating source matches were retained; confidence is based "
                    "on limited lexical evidence."
                ),
            )
        )

    if extra_caveats:
        caveats.extend(extra_caveats)

    metadata = {
        "scoring_policy": "v2_explainable",
        "confidence_band": confidence_band,
        "evidence_sentences": str(sentences_with_sources),
        "semantic_sentences": str(semantic_sentences),
        "fallback_sentences": str(fallback_sentences),
        "total_source_matches": str(total_source_matches),
        "source_group_count": str(len(source_groups)),
        "source_group_spans": str(total_group_spans),
    }

    if metadata_overrides:
        metadata.update(metadata_overrides)

    return ReportV2(
        source_groups=source_groups,
        match_groups=match_groups,
        caveats=caveats,
        metadata=metadata,
    )


async def collect_source_candidates(
    query: str,
    client: httpx.AsyncClient,
    connectors: list[SourceConnector] | None = None,
    per_source_caps: dict[str, int] | None = None,
    global_cap: int | None = None,
) -> list[NormalizedSourceCandidate]:
    active_connectors = (
        connectors if connectors is not None else get_source_connector_registry()
    )
    source_caps = per_source_caps if per_source_caps is not None else get_source_caps()
    max_total = (
        max(0, int(global_cap))
        if global_cap is not None
        else max(0, int(settings.PLAGIARISM_SOURCE_GLOBAL_MAX_CANDIDATES))
    )

    if max_total <= 0:
        return []

    connector_calls: list[tuple[SourceConnector, int]] = []
    for connector in active_connectors:
        if connector.name == "duckduckgo" and _duckduckgo_breaker_is_open():
            if _duckduckgo_should_log_breaker_open():
                print(
                    f"duckduckgo breaker open: skipping connector for {round(_duckduckgo_breaker_remaining_seconds(), 1)}s"
                )
            continue
        source_limit = max(0, int(source_caps.get(connector.name, max_total)))
        effective_limit = min(source_limit, max_total)
        if effective_limit <= 0:
            continue
        connector_calls.append((connector, effective_limit))

    if not connector_calls:
        return []

    connector_semaphore = asyncio.Semaphore(
        max(1, min(SOURCE_CONNECTOR_PARALLELISM, len(connector_calls)))
    )

    async def _run_connector(
        connector: SourceConnector,
        limit: int,
    ) -> list[NormalizedSourceCandidate]:
        async with connector_semaphore:
            try:
                raw_results = await connector.search(query, client, limit)
                if connector.name == "duckduckgo":
                    _duckduckgo_breaker_on_success()
            except Exception as e:
                if connector.name == "duckduckgo":
                    _duckduckgo_breaker_on_failure()
                    if _duckduckgo_should_log_error():
                        print(
                            f"{connector.name} search error: {_format_exception_summary(e)}"
                        )
                else:
                    print(
                        f"{connector.name} search error: {_format_exception_summary(e)}"
                    )
                return []

            normalized: list[NormalizedSourceCandidate] = []
            for raw in raw_results:
                try:
                    normalized.append(normalize_candidate_payload(raw))
                except Exception as e:
                    print(f"{connector.name} candidate parse error: {e}")
            return normalized

    per_connector_results = await asyncio.gather(
        *[
            _run_connector(connector=connector, limit=limit)
            for connector, limit in connector_calls
        ]
    )

    deduped_candidates: dict[int, NormalizedSourceCandidate] = {}
    key_to_group: dict[tuple[str, str], int] = {}
    group_counter = 0

    for connector_result in per_connector_results:
        for candidate in connector_result:
            candidate_keys = _dedupe_keys(candidate)
            matched_group = None
            for key in candidate_keys:
                existing_group = key_to_group.get(key)
                if existing_group is not None:
                    matched_group = existing_group
                    break

            if matched_group is None:
                matched_group = group_counter
                group_counter += 1
                deduped_candidates[matched_group] = candidate
            else:
                existing = deduped_candidates[matched_group]
                if _duplicate_quality_score(candidate) > _duplicate_quality_score(
                    existing
                ):
                    deduped_candidates[matched_group] = candidate

            for key in candidate_keys:
                key_to_group[key] = matched_group

    ranked = sorted(
        deduped_candidates.values(),
        key=lambda candidate: _rank_candidate(query, candidate),
        reverse=True,
    )
    return ranked[:max_total]


def get_random_user_agent() -> str:
    """Get a random user agent string."""
    return random.choice(USER_AGENTS)


def _apply_exclusion_pipeline(text: str) -> tuple[str, dict[str, int]]:
    stripped_text = text
    removed_quoted_segments = 0
    removed_parenthetical_citations = 0
    removed_numeric_citations = 0
    removed_reference_sections = 0

    stripped_text, quote_count = re.subn(r'"[^"]*"', "", stripped_text)
    removed_quoted_segments += quote_count
    stripped_text, quote_count = re.subn(r"'[^']*'", "", stripped_text)
    removed_quoted_segments += quote_count

    stripped_text, parenthetical_count = re.subn(
        r"\([^)]*\d{4}[^)]*\)",
        "",
        stripped_text,
    )
    removed_parenthetical_citations += parenthetical_count

    stripped_text, numeric_count = re.subn(
        r"\[(?:\d+\s*[,;-]\s*)*\d+\]",
        "",
        stripped_text,
    )
    removed_numeric_citations += numeric_count

    reference_heading_match = REFERENCE_SECTION_PATTERN.search(stripped_text)
    if reference_heading_match is None:
        reference_heading_match = REFERENCE_INLINE_PATTERN.search(stripped_text)
    if reference_heading_match:
        stripped_text = stripped_text[: reference_heading_match.start()]
        removed_reference_sections = 1

    stats = {
        "removed_quoted_segments": removed_quoted_segments,
        "removed_parenthetical_citations": removed_parenthetical_citations,
        "removed_numeric_citations": removed_numeric_citations,
        "removed_reference_sections": removed_reference_sections,
    }
    return stripped_text, stats


def _detect_ai_text(text: str) -> tuple[float, str]:
    if not text or len(text.strip()) < 50:
        return 0.0, "low"

    words = text.split()
    sentences = re.split(r"[.!?]+", text)
    sentences = [s.strip() for s in sentences if s.strip()]

    if not sentences:
        return 0.0, "low"

    score = 0.0

    avg_word_length = sum(len(w) for w in words) / len(words) if words else 0
    if avg_word_length > 6.5:
        score += 15

    sentence_lengths = [len(s.split()) for s in sentences if s.split()]
    if sentence_lengths:
        avg_sentence_len = sum(sentence_lengths) / len(sentence_lengths)
        if avg_sentence_len > 20:
            score += 20
        variance = sum((x - avg_sentence_len) ** 2 for x in sentence_lengths) / len(
            sentence_lengths
        )
        std_dev = variance**0.5
        if std_dev < 5:
            score += 15

    unique_words = len(set(words))
    total_words = len(words)
    if total_words > 0:
        lexical_diversity = unique_words / total_words
        if lexical_diversity > 0.7:
            score += 10
        elif lexical_diversity < 0.4:
            score += 20

    sentence_starters = [
        s.split()[0].lower() if s.split() else "" for s in sentences if s.split()
    ]
    unique_starters = len(set(sentence_starters))
    starter_diversity = (
        unique_starters / len(sentence_starters) if sentence_starters else 0
    )
    if starter_diversity < 0.3:
        score += 15

    transition_words = [
        "moreover",
        "furthermore",
        "additionally",
        "consequently",
        "therefore",
        "however",
        "nevertheless",
        "subsequently",
    ]
    transitions = sum(1 for w in words if w.lower() in transition_words)
    transition_ratio = transitions / total_words if total_words > 0 else 0
    if transition_ratio > 0.05:
        score += 10

    ai_indicators = [
        "it is important to note",
        "in conclusion",
        "research has shown",
        "studies have demonstrated",
        "it can be concluded",
        "based on the findings",
    ]
    for indicator in ai_indicators:
        if indicator.lower() in text.lower():
            score += 10
            break

    ai_score = min(100, max(0, score))

    if ai_score >= 70:
        confidence = "high"
    elif ai_score >= 40:
        confidence = "medium"
    else:
        confidence = "low"

    return ai_score, confidence


def _split_clean_sentences(text: str) -> list[str]:
    sentences = re.split(r"[.!?]+", text)
    return [
        sentence.strip()
        for sentence in sentences
        if sentence.strip() and len(sentence.strip()) > MIN_SENTENCE_LENGTH_CHARS
    ]


def _prepare_analysis_sentences(
    text: str,
    exclude_citations: bool,
) -> tuple[list[str], dict[str, str], list[ReportV2Caveat]]:
    input_chars = len(text)
    processed_text = text
    exclusion_stats = {
        "removed_quoted_segments": 0,
        "removed_parenthetical_citations": 0,
        "removed_numeric_citations": 0,
        "removed_reference_sections": 0,
    }

    if exclude_citations:
        processed_text, exclusion_stats = _apply_exclusion_pipeline(text)

    sentences = _split_clean_sentences(processed_text)
    remaining_chars = len(processed_text)
    removed_chars = max(0, input_chars - remaining_chars)
    removed_ratio = (removed_chars / input_chars) if input_chars > 0 else 0.0

    caveats: list[ReportV2Caveat] = []
    if exclude_citations and removed_chars > 0:
        caveats.append(
            ReportV2Caveat(
                code="EXCLUSION_CITATIONS_APPLIED",
                message=(
                    "Citation/reference exclusion removed quoted or reference-heavy text "
                    "before plagiarism analysis."
                ),
            )
        )

    if exclude_citations and removed_ratio >= HEAVY_EXCLUSION_RATIO_THRESHOLD:
        caveats.append(
            ReportV2Caveat(
                code="EXCLUSION_MAY_REDUCE_CONFIDENCE",
                message=(
                    "A large portion of text was excluded as citations/references; "
                    "plagiarism confidence may be reduced."
                ),
            )
        )

    if not sentences:
        caveats.append(
            ReportV2Caveat(
                code="INSUFFICIENT_ANALYZABLE_TEXT",
                message=(
                    "No analyzable sentences remained after preprocessing; returning "
                    "a minimal, non-failing plagiarism result."
                ),
            )
        )

    metadata = {
        "exclusion_requested": str(exclude_citations).lower(),
        "exclusion_applied": str(exclude_citations and removed_chars > 0).lower(),
        "excluded_characters": str(removed_chars),
        "excluded_characters_ratio": f"{removed_ratio:.4f}",
        "excluded_quoted_segments": str(exclusion_stats["removed_quoted_segments"]),
        "excluded_parenthetical_citations": str(
            exclusion_stats["removed_parenthetical_citations"]
        ),
        "excluded_numeric_citations": str(exclusion_stats["removed_numeric_citations"]),
        "excluded_reference_sections": str(
            exclusion_stats["removed_reference_sections"]
        ),
        "analyzable_sentences_before_cap": str(len(sentences)),
        "analyzable_text_minimal": str(not sentences).lower(),
    }

    return sentences, metadata, caveats


def extract_text_from_file(file_content: bytes, file_name: str) -> str:
    """
    Extract text from PDF, DOCX, or TXT file.

    Args:
        file_content: Raw file bytes
        file_name: Original filename to detect content type

    Returns:
        Extracted text content
    """
    if not file_content:
        raise ValueError("No file content provided")

    file_name_lower = file_name.lower()

    if file_name_lower.endswith(".pdf"):
        return _extract_text_from_pdf(file_content)
    elif file_name_lower.endswith(".docx"):
        return _extract_text_from_docx(file_content)
    elif file_name_lower.endswith(".txt"):
        return file_content.decode("utf-8", errors="ignore")
    else:
        raise ValueError(
            f"Unsupported file type: {file_name}. Supported: PDF, DOCX, TXT"
        )


def _extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF using PyMuPDF (fitz) - better text extraction with tables."""
    try:
        import fitz
    except ImportError:
        raise ImportError(
            "PyMuPDF (fitz) library is required for PDF extraction. Install with: pip install pymupdf"
        )

    doc = fitz.open(stream=file_content, filetype="pdf")
    text_parts = []

    for page_num, page in enumerate(doc):
        text = page.get_text("text")
        if text:
            text_parts.append(text)

        tables = page.find_tables()
        for table in tables:
            table_text = table.extract()
            if table_text:
                for row in table_text:
                    row_text = " | ".join(str(cell) if cell else "" for cell in row)
                    if row_text.strip():
                        text_parts.append(row_text)

    doc.close()
    return "\n".join(text_parts)


def _extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx library is required for DOCX extraction. Install with: pip install python-docx"
        )

    doc = Document(BytesIO(file_content))
    text_parts = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)

    return "\n".join(text_parts)


def split_into_sentences(text: str, exclude_citations: bool = False) -> List[str]:
    """
    Split text into sentences for analysis.

    Args:
        text: The input text
        exclude_citations: If True, remove quoted text before splitting

    Returns:
        List of sentences with length > 20 characters
    """
    sentences, _, _ = _prepare_analysis_sentences(text, exclude_citations)
    return sentences


def calculate_cosine_similarity(text1: str, text2: str) -> float:
    """
    Calculate cosine similarity between two texts using word frequency vectors.

    Returns:
        Similarity score between 0 and 1
    """
    words1 = text1.lower().split()
    words2 = text2.lower().split()

    if not words1 or not words2:
        return 0.0

    # Create vocabulary
    all_words = list(set(words1 + words2))

    # Create frequency vectors
    vector1 = [words1.count(word) for word in all_words]
    vector2 = [words2.count(word) for word in all_words]

    # Calculate dot product and magnitudes
    dot_product = sum(v1 * v2 for v1, v2 in zip(vector1, vector2))
    magnitude1 = math.sqrt(sum(v * v for v in vector1))
    magnitude2 = math.sqrt(sum(v * v for v in vector2))

    if magnitude1 == 0 or magnitude2 == 0:
        return 0.0

    return dot_product / (magnitude1 * magnitude2)


def calculate_ngram_similarity(text1: str, text2: str, n: int = 5) -> float:
    """
    Calculate n-gram similarity (Jaccard-like) between two texts.

    Args:
        text1: First text
        text2: Second text
        n: N-gram size (default 5 words)

    Returns:
        Similarity score between 0 and 1
    """

    def create_ngrams(text: str) -> Set[str]:
        # Clean and split into words
        words = re.sub(r"[^\w\s]", "", text.lower()).split()
        ngrams = set()
        for i in range(len(words) - n + 1):
            ngrams.add(" ".join(words[i : i + n]))
        return ngrams

    ngrams1 = create_ngrams(text1)
    ngrams2 = create_ngrams(text2)

    if not ngrams1 or not ngrams2:
        return 0.0

    # Count matches
    matches = len(ngrams1 & ngrams2)

    return matches / max(len(ngrams1), len(ngrams2))


def get_matching_ngrams(text1: str, text2: str, n: int = 5) -> Set[str]:
    words = re.sub(r"[^\w\s]", "", text1.lower()).split()
    ngrams1 = {
        " ".join(words[i : i + n]) for i in range(len(words) - n + 1) if len(words) >= n
    }

    words2 = re.sub(r"[^\w\s]", "", text2.lower()).split()
    ngrams2 = {
        " ".join(words2[i : i + n])
        for i in range(len(words2) - n + 1)
        if len(words2) >= n
    }

    return ngrams1 & ngrams2


# ============================================================
# W-SHINGLING WITH WINNOWING (Turnitin-style fingerprinting)
# ============================================================


def _hash_shingle(shingle: str) -> int:
    """Hash a shingle to an integer."""
    return hash(shingle)


def _winnow_hashes(hashes: list[int], window_size: int = 4) -> list[int]:
    """
    Winnowing: select minimum hash from each sliding window.
    Returns positions of selected hashes.
    """
    if len(hashes) < window_size:
        return list(range(len(hashes)))

    selected = []
    for i in range(len(hashes) - window_size + 1):
        window = hashes[i : i + window_size]
        min_val = min(window)
        min_pos = window.index(min_val)
        selected.append(i + min_pos)

    return sorted(set(selected))


def get_w_shingles(text: str, shingle_size: int = 7) -> Set[str]:
    """Generate w-shingles (7-10 word shingles)."""
    words = re.sub(r"[^\w\s]", "", text.lower()).split()
    if len(words) < shingle_size:
        return {text.lower()} if len(words) >= 3 else set()

    shingles = set()
    for i in range(len(words) - shingle_size + 1):
        shingle = " ".join(words[i : i + shingle_size])
        shingles.add(shingle)
    return shingles


def get_gapped_shingles(text: str, shingle_size: int = 7, gap: int = 1) -> Set[str]:
    """Generate gapped shingles (skip 1-2 tokens for modified copy detection)."""
    words = re.sub(r"[^\w\s]", "", text.lower()).split()
    if len(words) < shingle_size + gap:
        return set()

    shingles = set()
    for i in range(len(words) - shingle_size - gap + 1):
        # Take shingle_size words with gap tokens between them
        selected = [words[i + j + (j >= gap) * gap] for j in range(shingle_size)]
        shingle = " ".join(selected)
        shingles.add(shingle)
    return shingles


def calculate_w_similarity(text1: str, text2: str, shingle_size: int = 7) -> float:
    """Calculate w-shingle similarity with winnowing."""
    shingles1 = get_w_shingles(text1, shingle_size)
    shingles2 = get_w_shingles(text2, shingle_size)

    if not shingles1 or not shingles2:
        return 0.0

    # Hash and winnow
    hashes1 = sorted([_hash_shingle(s) for s in shingles1])
    hashes2 = sorted([_hash_shingle(s) for s in shingles2])

    winnow1 = set(_winnow_hashes(hashes1))
    winnow2 = set(_winnow_hashes(hashes2))

    if not winnow1 or not winnow2:
        return 0.0

    intersection = len(winnow1 & winnow2)
    return intersection / max(len(winnow1), len(winnow2))


# ============================================================
# PASSAGE ALIGNMENT (Turnitin-style continuous match detection)
# ============================================================


@dataclass
class MatchSpan:
    """A continuous matching passage."""

    start1: int
    end1: int
    start2: int
    end2: int
    text1: str
    text2: str
    similarity: float


def find_passage_alignments(
    text1: str, text2: str, min_match_words: int = 5
) -> List[MatchSpan]:
    """
    Find continuous matching passages between two texts.
    Uses token-level alignment to find runs of matching n-grams.
    """
    words1 = text1.lower().split()
    words2 = text2.lower().split()

    # Build n-gram index for text2
    n = min_match_words
    ngram_index = {}
    for i in range(len(words2) - n + 1):
        ngram = " ".join(words2[i : i + n])
        if ngram not in ngram_index:
            ngram_index[ngram] = []
        ngram_index[ngram].append(i)

    # Find matching n-gram positions
    matches = []
    for i in range(len(words1) - n + 1):
        ngram = " ".join(words1[i : i + n])
        if ngram in ngram_index:
            for j in ngram_index[ngram]:
                matches.append((i, j))

    if not matches:
        return []

    # Merge adjacent matches into spans
    matches.sort(key=lambda x: (x[0], x[1]))

    spans = []
    if matches:
        current_start1, current_start2 = matches[0]
        current_end1 = current_start1 + n
        current_end2 = current_start2 + n

        for i, (pos1, pos2) in enumerate(matches[1:]):
            # Check if this match is adjacent to current span
            if pos1 <= current_end1 + 2 and pos2 <= current_end2 + 2:
                # Extend current span
                current_end1 = pos1 + n
                current_end2 = pos2 + n
            else:
                # Save current span and start new one
                if current_end1 - current_start1 >= min_match_words:
                    spans.append(
                        MatchSpan(
                            start1=current_start1,
                            end1=current_end1,
                            start2=current_start2,
                            end2=current_end2,
                            text1=" ".join(words1[current_start1:current_end1]),
                            text2=" ".join(words2[current_start2:current_end2]),
                            similarity=1.0,
                        )
                    )
                current_start1, current_start2 = pos1, pos2
                current_end1 = pos1 + n
                current_end2 = pos2 + n

        # Don't forget the last span
        if current_end1 - current_start1 >= min_match_words:
            spans.append(
                MatchSpan(
                    start1=current_start1,
                    end1=current_end1,
                    start2=current_start2,
                    end2=current_end2,
                    text1=" ".join(words1[current_start1:current_end1]),
                    text2=" ".join(words2[current_start2:current_end2]),
                    similarity=1.0,
                )
            )

    return spans


# ============================================================
# MULTI-FEATURE PARAPHRASE CLASSIFIER
# ============================================================


def calculate_paraphrase_score(
    text1: str,
    text2: str,
    semantic_score: float,
    keyword_score: float,
) -> dict:
    """
    Multi-feature paraphrase detection classifier.

    Features:
    - Semantic cosine similarity (embeddings)
    - Token overlap ratio
    - Named entity overlap
    - Sentence length ratio
    - Rarity of terms (IDF-like)

    Returns:
        dict with paraphrase_likely (bool) and confidence (float)
    """
    words1 = set(re.sub(r"[^\w\s]", "", text1.lower()).split())
    words2 = set(re.sub(r"[^\w\s]", "", text2.lower()).split())

    # Token overlap
    intersection = words1 & words2
    union = words1 | words2
    token_overlap = len(intersection) / len(union) if union else 0

    # Length ratio
    len1, len2 = len(words1), len(words2)
    length_ratio = min(len1, len2) / max(len1, len2) if max(len1, len2) > 0 else 0

    # Common phrase detection (if too similar, likely copy not paraphrase)
    if token_overlap > 0.8:
        return {
            "paraphrase_likely": False,
            "confidence": 0.9,
            "reason": "high_token_overlap",
        }

    # Very common sentence (likely boilerplate)
    if len(words1) < 5:
        return {"paraphrase_likely": False, "confidence": 0.8, "reason": "too_short"}

    # Core logic: paraphrase when semantic is high but token overlap is moderate
    semantic_high = semantic_score > 0.6
    token_moderate = 0.3 < token_overlap < 0.7

    paraphrase_score = (
        (semantic_score * 0.5) + (token_overlap * 0.3) + (length_ratio * 0.2)
    )

    paraphrase_likely = semantic_high and token_moderate and paraphrase_score > 0.5
    confidence = min(0.95, paraphrase_score)

    return {
        "paraphrase_likely": paraphrase_likely,
        "confidence": confidence,
        "reason": "semantic_and_moderate_overlap"
        if paraphrase_likely
        else "insufficient_signals",
        "features": {
            "semantic_score": semantic_score,
            "token_overlap": token_overlap,
            "length_ratio": length_ratio,
            "paraphrase_score": paraphrase_score,
        },
    }


# ============================================================
# CROSS-ENCODER STYLE RE-RANKING (Simplified)
# ============================================================


def cross_encoder_score(text1: str, text2: str) -> float:
    """
    Simplified cross-encoder: direct pairwise comparison.
    In production, use a proper cross-encoder model.

    This combines multiple signals for "is this really similar?"
    """
    # Multiple similarity measures
    cosine = calculate_cosine_similarity(text1, text2)
    ngram_5 = calculate_ngram_similarity(text1, text2, 5)
    ngram_7 = calculate_ngram_similarity(text1, text2, 7)
    w_shingle = calculate_w_similarity(text1, text2, 7)

    # Passage alignment bonus
    alignments = find_passage_alignments(text1, text2, min_match_words=5)
    alignment_bonus = min(0.3, len(alignments) * 0.1)

    # Weighted combination (cross-encoder style)
    score = (
        cosine * 0.15
        + ngram_5 * 0.25
        + ngram_7 * 0.25
        + w_shingle * 0.25
        + alignment_bonus
    )

    return min(1.0, score)


def extract_quoted_text(text: str) -> list[tuple[int, int]]:
    """Extract positions of quoted text in the document."""
    positions = []
    for pattern in QUOTE_PATTERNS:
        for match in pattern.finditer(text):
            positions.append((match.start(), match.end()))
    return positions


def is_in_quoted_region(text: str, start: int, end: int) -> bool:
    """Check if a text range is inside quoted text."""
    quoted = extract_quoted_text(text)
    for q_start, q_end in quoted:
        if start >= q_start and end <= q_end:
            return True
    return False


def is_in_bibliography(text: str, start: int) -> bool:
    """Check if a position is in the bibliography section."""
    lines = text.split("\n")
    current_pos = 0
    in_bib = False

    for line in lines:
        line_end = current_pos + len(line)
        if current_pos <= start < line_end:
            return in_bib

        line_lower = line.lower().strip()
        for pattern in BIBLIOGRAPHY_PATTERNS:
            if pattern.search(line_lower):
                in_bib = True
                break

        current_pos = line_end + 1

    return in_bib


# ============================================================
# COMMON PHRASES & TEMPLATE TEXT EXCLUSIONS
# ============================================================

COMMON_ACADEMIC_PHRASES = {
    # English academic phrases
    "in this paper",
    "in this study",
    "in this research",
    "this paper discusses",
    "this paper examines",
    "this paper presents",
    "this paper proposes",
    "as shown in figure",
    "as illustrated in",
    "as can be seen from",
    "it is important to note",
    "it is worth noting",
    "it should be noted",
    "on the other hand",
    "on the contrary",
    "in conclusion",
    "in summary",
    "for example",
    "for instance",
    "in addition",
    "moreover",
    "furthermore",
    "however",
    "nevertheless",
    "nonetheless",
    "although",
    "whereas",
    "in terms of",
    "with regard to",
    "regarding",
    "concerning",
    "as mentioned above",
    "as stated previously",
    "as discussed earlier",
    "it is evident that",
    "it is clear that",
    "it can be concluded that",
    "the results show",
    "the findings indicate",
    "the data suggests",
    # Vietnamese academic phrases
    "trong nghiên cứu này",
    "trong bài báo này",
    "bài viết này trình bày",
    "như đã trình bày",
    "như thể hiện trong",
    "kết quả cho thấy",
    "cần lưu ý rằng",
    "điều quan trọng là",
    "tóm lại",
    "kết luận",
}

TEMPLATE_TEXT_PATTERNS = {
    # Course syllabus / policy templates
    r"course\s+(description|objectives|syllabus)",
    r"learning\s+outcomes",
    r"attendance\s+policy",
    r"academic\s+integrity",
    r"plagiarism\s+policy",
    r"grading\s+(policy|criteria|scale)",
    r"office\s+hours",
    r"prerequisite",
    r"textbook",
    r"course\s+materials",
    # Generic introduction templates
    r"^this\s+(paper|article|study|research)\s+is\s+(about|on)",
    r"the\s+purpose\s+of\s+(this|that)\s+(paper|study|research)",
    r"we\s+(will\s+)?discuss",
    r"this\s+chapter\s+(covers|presents|introduces)",
    # Footer/metadata templates
    r"all\s+rights\s+reserved",
    r"copyright\s+©",
    r"last\s+updated",
    r"contact\s+information",
}


def is_common_phrase(text: str) -> bool:
    """Check if text is a common academic phrase to exclude."""
    text_lower = text.lower().strip()
    words = text_lower.split()

    if len(words) <= 3:
        if text_lower in COMMON_ACADEMIC_PHRASES:
            return True

    phrase_window = " ".join(words[:5]) if len(words) >= 5 else text_lower
    for phrase in COMMON_ACADEMIC_PHRASES:
        if phrase in phrase_window or phrase_window in phrase:
            return True

    return False


def is_template_text(text: str) -> bool:
    """Detect template/syllabus/policy text."""
    text_lower = text.lower()

    for pattern in TEMPLATE_TEXT_PATTERNS:
        if re.search(pattern, text_lower, re.IGNORECASE):
            return True

    return False


def find_citation_markers(text: str, window_start: int, window_end: int) -> List[dict]:
    """Find citation markers near a text window."""
    window_text = text[max(0, window_start) : window_end]
    citations = []

    for pattern in CITATION_PATTERNS:
        for match in re.finditer(pattern, window_text):
            citations.append(
                {
                    "start": match.start(),
                    "end": match.end(),
                    "text": match.group(),
                }
            )

    return citations


def has_nearby_citation(
    text: str, match_start: int, match_end: int, window_sentences: int = 2
) -> bool:
    """Check if there's a citation within window of the match."""
    text_len = len(text)

    window_before_start = max(0, match_start - 500)
    window_after_end = min(text_len, match_end + 500)

    text_before = text[window_before_start:match_start]
    text_after = text[match_end:window_after_end]

    for pattern in CITATION_PATTERNS:
        if re.search(pattern, text_before):
            return True
        if re.search(pattern, text_after):
            return True

    return False


def merge_overlapping_spans(spans: List[tuple]) -> List[tuple]:
    """Merge overlapping/adjacent spans to avoid double-counting."""
    if not spans:
        return []

    sorted_spans = sorted(spans, key=lambda x: x[0])
    merged = [sorted_spans[0]]

    for start, end in sorted_spans[1:]:
        last_start, last_end = merged[-1]
        if start <= last_end + 1:
            merged[-1] = (last_start, max(last_end, end))
        else:
            merged.append((start, end))

    return merged


def calculate_confidence_and_type(
    passage_matches: list,
    semantic_similarity: int,
    keyword_similarity: int,
    source_credibility: int = 50,
) -> tuple[str, str]:
    """
    Calculate confidence score and match type.

    Returns:
        (confidence_level, match_type)
        confidence_level: high, medium, low
        match_type: exact, passage, semantic_only, possible_paraphrase
    """
    if not passage_matches:
        if semantic_similarity > keyword_similarity:
            return "low", "semantic_only"
        return "low", "possible_paraphrase"

    max_span_length = 0
    total_span_length = 0
    for pm in passage_matches:
        span_len = pm.get("end1", 0) - pm.get("start1", 0)
        max_span_length = max(max_span_length, span_len)
        total_span_length += span_len

    avg_span_length = total_span_length / len(passage_matches) if passage_matches else 0

    match_type = "passage"
    if max_span_length >= 20:
        match_type = "exact"
    elif max_span_length < 5:
        match_type = "possible_paraphrase"

    if semantic_similarity > keyword_similarity and max_span_length < 10:
        match_type = "possible_paraphrase"

    confidence = "medium"
    confidence_score = 0

    confidence_score += min(40, max_span_length)
    confidence_score += min(30, avg_span_length)
    confidence_score += source_credibility * 0.3

    if match_type == "exact":
        confidence_score += 20
    elif match_type == "passage":
        confidence_score += 10

    if confidence_score >= 60:
        confidence = "high"
    elif confidence_score >= 30:
        confidence = "medium"
    else:
        confidence = "low"

    return confidence, match_type


def calculate_coverage_score(
    text: str,
    results: list,
    exclude_quotes: bool = True,
    exclude_bibliography: bool = True,
    exclude_common: bool = True,
    exclude_template: bool = True,
    min_word_threshold: int = 10,
    citation_reduction: bool = True,
) -> dict:
    """
    Calculate coverage-based similarity score (Turnitin-style).
    Uses matched character/token spans instead of simple stem counting.
    """
    all_words = re.sub(r"[^\w\s]", "", text.lower()).split()
    total_words = len(all_words)
    total_chars = len(text)

    if total_words == 0:
        return {
            "overall_score": 0,
            "matched_words": 0,
            "total_words": 0,
            "excluded_words": 0,
            "coverage_percent": 0,
            "unique_matched_chars": 0,
            "total_chars": total_chars,
        }

    matched_spans: List[tuple] = []
    excluded_words = 0
    citation_reduced_count = 0

    for result in results:
        sentence = result.sentence
        sentence_words = re.sub(r"[^\w\s]", "", sentence.lower()).split()
        sentence_start = text.find(sentence)

        if sentence_start == -1:
            sentence_start = 0

        is_quoted = exclude_quotes and is_in_quoted_region(
            text, sentence_start, sentence_start + len(sentence)
        )
        is_bib = exclude_bibliography and is_in_bibliography(text, sentence_start)

        if is_quoted or is_bib:
            excluded_words += len(sentence_words)
            continue

        if exclude_common and is_common_phrase(sentence):
            excluded_words += len(sentence_words)
            continue

        if exclude_template and is_template_text(sentence):
            excluded_words += len(sentence_words)
            continue

        if len(sentence_words) < min_word_threshold:
            excluded_words += len(sentence_words)
            continue

        sources = getattr(result, "sources", []) or []
        has_matches = any(s.similarity > 0 for s in sources)

        if not has_matches:
            continue

        match_start = sentence_start
        match_end = sentence_start + len(sentence)

        if citation_reduction and has_nearby_citation(text, match_start, match_end):
            citation_reduced_count += 1
            continue

        matched_spans.append((match_start, match_end))

    merged_spans = merge_overlapping_spans(matched_spans)

    matched_chars = sum(end - start for start, end in merged_spans)
    coverage_percent = (matched_chars / total_chars * 100) if total_chars > 0 else 0

    analyzable_words = total_words - excluded_words
    if analyzable_words <= 0:
        return {
            "overall_score": 0,
            "matched_words": 0,
            "total_words": analyzable_words,
            "excluded_words": excluded_words,
            "coverage_percent": 0,
            "unique_matched_chars": matched_chars,
            "total_chars": total_chars,
            "merged_spans_count": len(merged_spans),
            "citation_reduced_count": citation_reduced_count,
        }

    overall_score = min(100, round(coverage_percent))

    return {
        "overall_score": overall_score,
        "matched_words": len(matched_spans),
        "total_words": analyzable_words,
        "excluded_words": excluded_words,
        "coverage_percent": round(coverage_percent, 2),
        "unique_matched_chars": matched_chars,
        "total_chars": total_chars,
        "merged_spans_count": len(merged_spans),
        "citation_reduced_count": citation_reduced_count,
    }


def calculate_turnitin_score(
    text: str,
    results: list,
    exclude_quotes: bool = True,
    exclude_bibliography: bool = True,
    small_match_threshold: int = 10,
) -> dict:
    all_words = re.sub(r"[^\w\s]", "", text.lower()).split()
    total_words = len(all_words)

    if total_words == 0:
        return {
            "overall_score": 0,
            "matched_words": 0,
            "total_words": 0,
            "excluded_words": 0,
        }

    matched_stems_set: set[str] = set()
    excluded_words = 0

    all_stems = stem_text(text)

    for idx, result in enumerate(results):
        sentence = result.sentence
        sentence_words = re.sub(r"[^\w\s]", "", sentence.lower()).split()
        sentence_start = text.find(sentence)

        is_quoted = exclude_quotes and is_in_quoted_region(
            text, sentence_start, sentence_start + len(sentence)
        )
        is_bib = exclude_bibliography and is_in_bibliography(text, sentence_start)

        if is_quoted or is_bib:
            excluded_words += len(sentence_words)
            continue

        if len(sentence_words) < small_match_threshold:
            excluded_words += len(sentence_words)
            continue

        matched_ngrams = getattr(result, "matched_ngrams", [])
        if matched_ngrams:
            for ngram in matched_ngrams:
                ngram_words = ngram.split()
                for word in ngram_words:
                    if word and len(word) > 1:
                        matched_stems_set.add(stem_word(word))
        else:
            if result.sources and any(s.similarity > 0 for s in result.sources):
                for word in sentence_words:
                    if len(word) > 1:
                        matched_stems_set.add(stem_word(word))

    unique_matched_words = len(matched_stems_set & all_stems)
    analyzable_words = total_words - excluded_words

    if analyzable_words <= 0:
        return {
            "overall_score": 0,
            "matched_words": unique_matched_words,
            "total_words": total_words,
            "excluded_words": excluded_words,
        }

    overall_score = min(100, round((unique_matched_words / analyzable_words) * 100))

    return {
        "overall_score": overall_score,
        "matched_words": unique_matched_words,
        "total_words": analyzable_words,
        "excluded_words": excluded_words,
    }


async def search_duckduckgo(query: str, client: httpx.AsyncClient) -> List[str]:
    """
    Search DuckDuckGo HTML version for URLs matching the query.

    Returns:
        List of URLs found (max 8)
    """
    urls = []
    try:
        connector = DuckDuckGoConnector()
        candidates = await connector.search(
            query,
            client,
            max(0, int(settings.PLAGIARISM_SOURCE_DUCKDUCKGO_MAX_CANDIDATES)),
        )
        urls = [candidate.canonical_url for candidate in candidates]
    except Exception as e:
        print(f"DuckDuckGo search error: {_format_exception_summary(e)}")

    return urls


async def search_crossref(query: str, client: httpx.AsyncClient) -> List[str]:
    """
    Search CrossRef API for academic paper URLs.

    Returns:
        List of DOI URLs found (max 5)
    """
    urls = []
    try:
        connector = CrossRefConnector()
        candidates = await connector.search(
            query,
            client,
            max(0, int(settings.PLAGIARISM_SOURCE_CROSSREF_MAX_CANDIDATES)),
        )
        urls = [candidate.canonical_url for candidate in candidates]
    except Exception as e:
        print(f"CrossRef search error: {e}")

    return urls


# Stopwords for query rewriting
STOPWORDS = {
    "a",
    "an",
    "the",
    "is",
    "are",
    "was",
    "were",
    "be",
    "been",
    "being",
    "have",
    "has",
    "had",
    "do",
    "does",
    "did",
    "will",
    "would",
    "could",
    "should",
    "may",
    "might",
    "must",
    "shall",
    "can",
    "need",
    "dare",
    "to",
    "of",
    "in",
    "for",
    "on",
    "with",
    "at",
    "by",
    "from",
    "as",
    "into",
    "through",
    "during",
    "before",
    "after",
    "above",
    "below",
    "between",
    "under",
    "again",
    "further",
    "then",
    "once",
    "here",
    "there",
    "when",
    "where",
    "why",
    "how",
    "all",
    "each",
    "few",
    "more",
    "most",
    "other",
    "some",
    "such",
    "no",
    "nor",
    "not",
    "only",
    "own",
    "same",
    "so",
    "than",
    "too",
    "very",
    "just",
    "also",
    "now",
    "and",
    "but",
    "or",
    "yet",
    "if",
    "because",
    "although",
    "while",
    "that",
    "which",
    "who",
    "whom",
    "this",
    "these",
    "those",
    "it",
    "its",
    "they",
    "them",
    "their",
    "we",
    "you",
    "your",
    "he",
    "she",
    "him",
    "her",
    "his",
    "i",
    "me",
    "my",
    "us",
    "our",
    # Vietnamese stopwords
    "và",
    "của",
    "trong",
    "được",
    "với",
    "là",
    "các",
    "có",
    "không",
    "để",
    "từ",
    "về",
    "cho",
    "này",
    "đó",
    "những",
    "một",
    "vì",
    "nên",
    "khi",
    "nếu",
    "hoặc",
    "vẫn",
    "đã",
    "đang",
    "sẽ",
    "phải",
}


def _remove_stopwords(text: str) -> str:
    """Remove stopwords from text."""
    words = text.lower().split()
    return " ".join(w for w in words if w not in STOPWORDS)


def _extract_noun_phrases(text: str) -> str:
    """Extract noun phrases (simple heuristic: 2-4 word sequences of non-stopwords."""
    words = text.lower().split()
    phrases = []
    for n in range(2, min(5, len(words) + 1)):
        for i in range(len(words) - n + 1):
            phrase_words = words[i : i + n]
            # Keep if majority are not stopwords
            non_stop = sum(1 for w in phrase_words if w not in STOPWORDS)
            if non_stop >= n - 1:
                phrases.append(" ".join(phrase_words))
    # Return top phrases by frequency
    from collections import Counter

    phrase_counts = Counter(phrases)
    top_phrases = [p for p, _ in phrase_counts.most_common(5)]
    return " ".join(top_phrases)


def _extract_key_ngrams(text: str, n: int = 4) -> str:
    """Extract key n-grams (4-word phrases) as unique identifiers."""
    words = text.lower().split()
    if len(words) < n:
        return text.lower()
    ngrams = [" ".join(words[i : i + n]) for i in range(len(words) - n + 1)]
    return " ".join(ngrams[:3])  # Top 3 n-grams


def _generate_query_variants(text: str) -> list[str]:
    """Generate 2-4 query variants for better search results."""
    variants = [
        text,  # Original
        _remove_stopwords(text),  # Without stopwords
    ]

    # Add noun phrases if text is long enough
    if len(text.split()) > 5:
        noun_phrases = _extract_noun_phrases(text)
        if noun_phrases:
            variants.append(noun_phrases)

    # Add key n-grams
    key_ngrams = _extract_key_ngrams(text, 4)
    if key_ngrams and key_ngrams != text.lower():
        variants.append(key_ngrams)

    # Remove duplicates while preserving order
    seen = set()
    unique_variants = []
    for v in variants:
        v_clean = v.strip()
        if v_clean and v_clean not in seen:
            seen.add(v_clean)
            unique_variants.append(v_clean)

    return unique_variants[:4]  # Max 4 variants


def _create_chunks(
    sentences: list[str], chunk_size: int = 5, overlap: int = 2
) -> list[tuple[int, int, str]]:
    """
    Create chunks of sentences with overlap.

    Returns:
        List of (start_idx, end_idx, chunk_text) tuples
    """
    chunks = []
    n = len(sentences)

    for i in range(0, n, chunk_size - overlap):
        end = min(i + chunk_size, n)
        chunk_sentences = sentences[i:end]
        chunk_text = " ".join(chunk_sentences)
        chunks.append((i, end, chunk_text))

        if end >= n:
            break

    return chunks


async def search_web(query: str, client: httpx.AsyncClient) -> List[str]:
    """
    Search with query variants for better source retrieval.

    Returns:
        Combined list of unique URLs (max 10)
    """
    # Generate query variants
    query_variants = _generate_query_variants(query)

    all_candidates = []
    seen_urls = set()

    # Search with each variant
    for variant in query_variants[:3]:  # Max 3 variants
        if len(variant) < 10:  # Skip very short queries
            continue

        try:
            candidates = await collect_source_candidates(query=variant, client=client)

            # Add unique candidates
            for candidate in candidates:
                url = candidate.canonical_url
                if url not in seen_urls:
                    seen_urls.add(url)
                    all_candidates.append(candidate)

            if len(all_candidates) >= 10:
                break
        except Exception as e:
            print(f"Search error for variant '{variant}': {e}")
            continue

    # Cache the combined results (use original query as key)
    _cache_query_candidates(query=query, candidates=all_candidates)

    return [candidate.canonical_url for candidate in all_candidates[:10]]


# Boilerplate patterns for removal
BOILERPLATE_TAGS = [
    "script",
    "style",
    "nav",
    "header",
    "footer",
    "aside",
    "noscript",
    "iframe",
    "form",
    "button",
    "input",
    "select",
    "textarea",
]

BOILERPLATE_CLASSES = [
    "menu",
    "nav",
    "navigation",
    "navbar",
    "sidebar",
    "widget",
    "footer",
    "header",
    "banner",
    "cookie",
    "popup",
    "modal",
    "advertisement",
    "ad-",
    "social",
    "share",
    "comment",
    "related",
    "sidebar",
    "breadcrumb",
    "pagination",
    "copyright",
    "sitemap",
    "login",
    "register",
    "signup",
    "newsletter",
    "subscribe",
]

BOILERPLATE_IDS = [
    "menu",
    "nav",
    "navigation",
    "sidebar",
    "footer",
    "header",
    "cookie",
    "popup",
    "modal",
    "advertisement",
    "sidebar",
]


def _remove_boilerplate(soup: BeautifulSoup) -> BeautifulSoup:
    """Remove boilerplate content from HTML (menu, footer, cookies, etc.)"""
    try:
        for tag in soup(BOILERPLATE_TAGS):
            tag.decompose()
    except Exception:
        pass

    try:
        for element in soup.find_all(class_=True):
            if element is None:
                continue
            class_list = element.get("class") or []
            class_name = " ".join(class_list) if class_list else ""
            if class_name and any(
                bc in class_name.lower() for bc in BOILERPLATE_CLASSES
            ):
                element.decompose()
    except Exception:
        pass

    try:
        for element in soup.find_all(id=True):
            if element is None:
                continue
            element_id = element.get("id") or ""
            if element_id and any(bc in element_id.lower() for bc in BOILERPLATE_IDS):
                element.decompose()
    except Exception:
        pass

    try:
        for element in soup.find_all(
            style=lambda x: x and "display:none" in x.replace(" ", "")
        ):
            if element:
                element.decompose()
    except Exception:
        pass

    try:
        boilerplate_texts = [
            "cookie",
            "privacy policy",
            "terms of service",
            "copyright",
            "all rights reserved",
            "powered by",
            "subscribe to",
            "newsletter",
            "follow us",
            "share this",
            "last updated",
        ]
        for element in soup.find_all(string=True):
            if element and element.string:
                text = element.string.lower().strip()
                if any(bt in text for bt in boilerplate_texts):
                    parent = element.find_parent()
                    if parent:
                        parent.decompose()
    except Exception:
        pass

    return soup


def _normalize_text(text: str) -> str:
    """Normalize unicode, quotes, hyphenation, and line breaks."""
    # Normalize unicode quotes to ASCII
    text = text.replace("\u201c", '"').replace("\u201d", '"')
    text = text.replace("\u2018", "'").replace("\u2019", "'")
    text = text.replace("\u2013", "-").replace("\u2014", "--")
    text = text.replace("\u2026", "...")
    text = text.replace("\u00a0", " ")

    # Fix hyphenation (line break in middle of word)
    text = re.sub(r"(\w+)-\s+(\w+)", r"\1\2", text)

    # Normalize line breaks
    text = re.sub(r"\n+", "\n", text)
    text = re.sub(r" {2,}", " ", text)

    # Clean quotes
    text = re.sub(r'"(\w)"', r'"\1"', text)
    text = re.sub(r"'(\w)'", r"'\1'", text)

    # Remove multiple spaces
    text = re.sub(r" +", " ", text)

    return text.strip()


def _extract_main_content(soup: BeautifulSoup) -> str:
    """Extract main content using heuristics (article, main, content)"""
    try:
        main_tags = soup.find_all(["article", "main", "div"])
    except Exception:
        return ""

    content_texts = []

    for tag in main_tags:
        if tag is None:
            continue
        try:
            tag_id = (tag.get("id") or "").lower()
            tag_class = " ".join(tag.get("class") or []).lower()
        except Exception:
            continue

        if any(x in tag_id for x in ["nav", "menu", "sidebar", "footer", "header"]):
            continue
        if any(x in tag_class for x in ["nav", "menu", "sidebar", "footer", "header"]):
            continue

        try:
            text = tag.get_text(separator=" ", strip=True)
            if len(text) > 100:
                content_texts.append((len(text), text))
        except Exception:
            continue

    if content_texts:
        content_texts.sort(reverse=True)
        return content_texts[0][1]

    return ""


async def fetch_page_content(url: str, client: httpx.AsyncClient) -> str:
    """
    Fetch and clean text content from a URL.

    Returns:
        Cleaned text content (max 5000 chars)
    """
    try:
        response = await client.get(
            url,
            headers={
                "User-Agent": get_random_user_agent(),
                "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
                "Accept-Language": "en-US,en;q=0.5",
            },
            timeout=8.0,
            follow_redirects=True,
        )

        if response.status_code != 200:
            return ""

        # Parse HTML with BeautifulSoup
        soup = BeautifulSoup(response.text, "html.parser")

        # Remove boilerplate (menu, footer, cookies, etc.)
        soup = _remove_boilerplate(soup)

        # Try to extract main content
        text = _extract_main_content(soup)
        if not text:
            # Fallback to getting all text
            text = soup.get_text(separator=" ", strip=True)

        # Normalize text (unicode, quotes, hyphenation)
        text = _normalize_text(text)

        return text[:5000]
    except Exception as e:
        print(f"Error fetching {url}: {e}")
        return ""


async def check_sentence(
    sentence: str,
    client: httpx.AsyncClient,
    semaphore: asyncio.Semaphore,
    use_ai_similarity: bool,
    user_key: str,
) -> SentenceResult:
    """
    Check a single sentence for plagiarism.

    Returns:
        SentenceResult with similarity scores and sources
    """
    async with semaphore:
        # Add random delay to avoid rate limiting
        await asyncio.sleep(random.uniform(0.5, 1.5))

        urls = await search_web(sentence, client)
        ranked_candidates = _get_cached_query_candidates(sentence)
        if ranked_candidates is None:
            ranked_candidates = [
                NormalizedSourceCandidate(
                    source="legacy",
                    canonical_url=url,
                    title=url,
                )
                for url in urls
            ]

        max_similarity = 0.0
        semantic_similarity = 0
        used_ai = False
        fallback_used = False
        analysis_method = "keyword"
        best_content = ""
        matched_sources: List[SourceMatch] = []
        max_keyword_similarity = 0.0

        top_candidates = ranked_candidates[:SIMILARITY_TOP_K_CANDIDATES]
        for candidate in top_candidates:
            content = await fetch_page_content(candidate.canonical_url, client)
            comparison_text = (
                content
                if content and len(content) > 100
                else (candidate.snippet or "").strip()
            )
            if not comparison_text:
                continue

            cosine_sim = calculate_cosine_similarity(sentence, comparison_text)
            ngram_sim = calculate_ngram_similarity(sentence, comparison_text, 5)
            keyword_sim = max(cosine_sim, ngram_sim)

            if keyword_sim > max_keyword_similarity:
                max_keyword_similarity = keyword_sim

            similarity = max(keyword_sim, max_keyword_similarity)

            if similarity > max_similarity:
                max_similarity = similarity
                best_content = comparison_text

            if similarity > 0.15:
                matching_ngrams = get_matching_ngrams(sentence, comparison_text, 5)

                # Find passage-level alignments (Turnitin-style)
                passage_matches = []
                if len(sentence.split()) >= 5 and len(comparison_text.split()) >= 5:
                    try:
                        spans = find_passage_alignments(
                            sentence, comparison_text, min_match_words=5
                        )
                        for span in spans:
                            passage_matches.append(
                                {
                                    "text1": span.text1,
                                    "text2": span.text2,
                                    "start1": span.start1,
                                    "end1": span.end1,
                                    "start2": span.start2,
                                    "end2": span.end2,
                                    "similarity": round(span.similarity * 100),
                                }
                            )
                    except Exception:
                        pass  # Skip passage alignment if it fails

                matched_sources.append(
                    SourceMatch(
                        url=candidate.canonical_url,
                        similarity=round(similarity * 100),
                        matched_ngrams=list(matching_ngrams),
                        passage_matches=passage_matches,
                    )
                )

        if best_content:
            semantic_result = await calculate_semantic_similarity(
                sentence,
                best_content,
                use_ai=use_ai_similarity,
                user_key=user_key,
            )
            semantic_similarity = round(semantic_result.similarity * 100)
            used_ai = semantic_result.used_ai
            fallback_used = semantic_result.fallback_used
            analysis_method = semantic_result.method
            max_similarity = max(max_similarity, semantic_result.similarity)

        paraphrase_detected = semantic_similarity > 0 and semantic_similarity > (
            max_keyword_similarity * 100
        )

        for source in matched_sources:
            confidence, match_type = calculate_confidence_and_type(
                source.passage_matches,
                semantic_similarity=semantic_similarity,
                keyword_similarity=source.similarity,
                source_credibility=50,
            )
            source.confidence_score = confidence
            source.match_type = match_type

        matched_sources.sort(key=lambda x: x.similarity, reverse=True)

        all_ngrams = []
        for source in matched_sources:
            all_ngrams.extend(source.matched_ngrams)

        return SentenceResult(
            sentence=sentence,
            similarity=round(max_similarity * 100),
            semantic_similarity=semantic_similarity,
            sources=matched_sources[:5],
            matched_ngrams=list(set(all_ngrams)),
            is_plagiarized=max_similarity > 0.5,
            used_ai=used_ai,
            fallback_used=fallback_used,
            analysis_method=analysis_method,
            paraphrase_detected=paraphrase_detected,
        )


async def check_plagiarism(
    request: PlagiarismCheckRequest,
    *,
    user_key: str | None = None,
) -> PlagiarismCheckResponse:
    """
    Main function to check text for plagiarism.

    Args:
        request: The plagiarism check request

    Returns:
        PlagiarismCheckResponse with overall scores and detailed results
    """
    resolved_user_key = user_key or "anon:global"

    sentences, exclusion_metadata, exclusion_caveats = _prepare_analysis_sentences(
        request.text,
        request.exclude_citations,
    )

    if not sentences:
        quota_info = get_quota_info(user_key=resolved_user_key)
        ai_detection_score, ai_detection_confidence = _detect_ai_text(request.text)
        return PlagiarismCheckResponse(
            overall_score=0,
            plagiarism_percentage=0,
            total_sentences=0,
            plagiarized_sentences=0,
            results=[],
            used_ai_similarity=False,
            fallback_used=False,
            analysis_method="keyword",
            ai_quota_remaining=quota_info["remaining"],
            ai_quota_percent=quota_info["usage_percent"],
            source_counts=None,
            source_failures=None,
            quota_mode=quota_info.get("quota_mode"),
            ai_detection_score=ai_detection_score,
            ai_detection_confidence=ai_detection_confidence,
            report_v2=_build_report_v2(
                [],
                metadata_overrides=exclusion_metadata,
                extra_caveats=exclusion_caveats,
            ),
        )

    # Limit number of sentences
    sentences = sentences[: request.max_sentences]
    exclusion_metadata["analyzable_sentences_after_cap"] = str(len(sentences))

    # Create semaphore for rate limiting
    semaphore = asyncio.Semaphore(MAX_CONCURRENT_REQUESTS)
    semantic_budget = max(0, int(settings.SEMANTIC_MAX_CHECKS_PER_REQUEST))

    # Check all sentences concurrently (with rate limiting)
    async with httpx.AsyncClient() as client:
        tasks = [
            check_sentence(
                sentence,
                client,
                semaphore,
                use_ai_similarity=(request.use_ai_similarity and idx < semantic_budget),
                user_key=resolved_user_key,
            )
            for idx, sentence in enumerate(sentences)
        ]
        results = await asyncio.gather(*tasks)

    # Calculate overall statistics using coverage-based scoring (Turnitin-style)
    turnitin_stats = calculate_coverage_score(
        request.text,
        results,
        exclude_quotes=True,
        exclude_bibliography=request.exclude_bibliography,
        exclude_common=getattr(request, "exclude_common_phrases", True),
        exclude_template=getattr(request, "exclude_template_text", True),
        min_word_threshold=getattr(request, "min_word_threshold", 10),
        citation_reduction=getattr(request, "citation_severity_reduction", True),
    )

    # Fallback to legacy stem-count if coverage returns 0 but sources exist
    if turnitin_stats["overall_score"] == 0 and any(r.sources for r in results):
        turnitin_stats = calculate_turnitin_score(
            request.text,
            results,
            exclude_quotes=True,
            exclude_bibliography=request.exclude_bibliography,
            small_match_threshold=getattr(request, "min_word_threshold", 10),
        )

    overall_score = turnitin_stats["overall_score"]
    plagiarism_percentage = turnitin_stats["overall_score"]

    plagiarized_count = sum(1 for r in results if r.is_plagiarized)

    used_ai_similarity = any(r.used_ai for r in results)
    fallback_used = any(r.fallback_used for r in results)
    analysis_method = (
        "hybrid"
        if used_ai_similarity and fallback_used
        else "semantic"
        if used_ai_similarity
        else "keyword"
    )

    filtered_results = results
    small_matches_removed = 0

    if request.exclude_small_matches > 0:
        filtered_results, small_matches_removed = _filter_small_matches(
            filtered_results, request.exclude_small_matches
        )
        exclusion_metadata["small_matches_removed"] = str(small_matches_removed)
        exclusion_metadata["small_match_threshold"] = str(request.exclude_small_matches)

    if request.exclude_small_sources:
        before_filter = len(filtered_results)
        filtered_results = _filter_small_sources(filtered_results, min_source_count=3)
        sources_removed = before_filter - len(filtered_results)
        exclusion_metadata["small_sources_removed"] = str(sources_removed)

    source_type_filter = getattr(request, "source_type_filter", None)
    if source_type_filter and len(source_type_filter) > 0:
        filtered_results = _filter_by_source_type(filtered_results, source_type_filter)
        exclusion_metadata["source_type_filter"] = ",".join(source_type_filter)

    contribution_threshold = getattr(request, "source_contribution_threshold", 0)
    if contribution_threshold > 0:
        filtered_results = _filter_by_contribution(
            filtered_results, contribution_threshold
        )
        exclusion_metadata["contribution_threshold"] = str(contribution_threshold)

    recalculate_for_response = len(filtered_results) > 0

    if recalculate_for_response:
        turnitin_stats_filtered = calculate_turnitin_score(
            request.text,
            filtered_results,
            exclude_quotes=True,
            exclude_bibliography=request.exclude_bibliography,
            small_match_threshold=10,
        )
        overall_score = turnitin_stats_filtered["overall_score"]
        plagiarism_percentage = turnitin_stats_filtered["overall_score"]
        plagiarized_count = sum(1 for r in filtered_results if r.is_plagiarized)
    else:
        overall_score = 0
        plagiarism_percentage = 0
        plagiarized_count = 0

    quota_info = get_quota_info(user_key=resolved_user_key)
    source_counts = (
        _build_source_counts(filtered_results) if recalculate_for_response else {}
    )

    if small_matches_removed > 0:
        exclusion_caveats.append(
            ReportV2Caveat(
                code="SMALL_MATCHES_FILTERED",
                message=(
                    f"{small_matches_removed} matches smaller than {request.exclude_small_matches} "
                    "words were excluded from analysis."
                ),
            )
        )

    ai_detection_score, ai_detection_confidence = _detect_ai_text(request.text)

    return PlagiarismCheckResponse(
        overall_score=overall_score,
        plagiarism_percentage=plagiarism_percentage,
        total_sentences=len(filtered_results),
        plagiarized_sentences=plagiarized_count,
        results=filtered_results,
        used_ai_similarity=used_ai_similarity,
        fallback_used=fallback_used,
        analysis_method=analysis_method,
        ai_quota_remaining=quota_info["remaining"],
        ai_quota_percent=quota_info["usage_percent"],
        source_counts=source_counts,
        source_failures=None,
        quota_mode=quota_info.get("quota_mode"),
        ai_detection_score=ai_detection_score,
        ai_detection_confidence=ai_detection_confidence,
        report_v2=_build_report_v2(
            filtered_results,
            metadata_overrides=exclusion_metadata,
            extra_caveats=exclusion_caveats,
        ),
    )


async def check_plagiarism_streaming(
    request: PlagiarismCheckRequest,
    *,
    user_key: str | None = None,
):
    resolved_user_key = user_key or "anon:global"

    debug_metadata = {
        "candidates_fetched": 0,
        "sources_parsed": 0,
        "spans_found": 0,
        "sentences_processed": 0,
    }

    yield {
        "progress": 0,
        "current": 0,
        "total": 0,
        "status": "preparing",
        "message": "Preparing text for analysis...",
        "stage": "preparing",
        "debug": debug_metadata,
    }

    sentences, exclusion_metadata, exclusion_caveats = _prepare_analysis_sentences(
        request.text,
        request.exclude_citations,
    )

    if not sentences:
        quota_info = get_quota_info(user_key=resolved_user_key)
        ai_detection_score, ai_detection_confidence = _detect_ai_text(request.text)
        yield {
            "progress": 100,
            "current": 0,
            "total": 0,
            "status": "complete",
            "message": "Analysis complete",
        }
        yield PlagiarismCheckResponse(
            overall_score=0,
            plagiarism_percentage=0,
            total_sentences=0,
            plagiarized_sentences=0,
            results=[],
            used_ai_similarity=False,
            fallback_used=False,
            analysis_method="keyword",
            ai_quota_remaining=quota_info["remaining"],
            ai_quota_percent=quota_info["usage_percent"],
            source_counts=None,
            source_failures=None,
            quota_mode=quota_info.get("quota_mode"),
            ai_detection_score=ai_detection_score,
            ai_detection_confidence=ai_detection_confidence,
            report_v2=_build_report_v2(
                [],
                metadata_overrides=exclusion_metadata,
                extra_caveats=exclusion_caveats,
            ),
        ).model_dump()
        return

    sentences = sentences[: request.max_sentences]
    total_sentences = len(sentences)
    exclusion_metadata["analyzable_sentences_after_cap"] = str(total_sentences)

    yield {
        "progress": 5,
        "current": 0,
        "total": total_sentences,
        "status": "retrieval",
        "message": f"Retrieving candidates for {total_sentences} sentences...",
        "stage": "retrieval",
        "debug": debug_metadata,
    }

    semaphore = asyncio.Semaphore(MAX_CONCURRENT_REQUESTS)
    semantic_budget = max(0, int(settings.SEMANTIC_MAX_CHECKS_PER_REQUEST))

    results = []
    processed_count = 0

    async with httpx.AsyncClient() as client:
        for idx, sentence in enumerate(sentences):
            result = await check_sentence(
                sentence,
                client,
                semaphore,
                use_ai_similarity=(request.use_ai_similarity and idx < semantic_budget),
                user_key=resolved_user_key,
            )
            results.append(result)
            processed_count += 1

            debug_metadata["sentences_processed"] = processed_count
            if result.sources:
                debug_metadata["sources_parsed"] += len(result.sources)
                for src in result.sources:
                    if src.passage_matches:
                        debug_metadata["spans_found"] += len(src.passage_matches)

            current_stage = (
                "retrieval"
                if idx < total_sentences * 0.3
                else "download"
                if idx < total_sentences * 0.6
                else "align"
            )

            progress = int((processed_count / total_sentences) * 80) + 5
            yield {
                "progress": progress,
                "current": processed_count,
                "total": total_sentences,
                "status": current_stage,
                "message": f"Processing sentence {processed_count}/{total_sentences}",
                "stage": current_stage,
                "debug": debug_metadata,
            }

    yield {
        "progress": 85,
        "current": total_sentences,
        "total": total_sentences,
        "status": "scoring",
        "message": "Calculating similarity scores...",
        "stage": "rerank",
        "debug": debug_metadata,
    }

    turnitin_stats = calculate_coverage_score(
        request.text,
        results,
        exclude_quotes=True,
        exclude_bibliography=request.exclude_bibliography,
        exclude_common=getattr(request, "exclude_common_phrases", True),
        exclude_template=getattr(request, "exclude_template_text", True),
        min_word_threshold=getattr(request, "min_word_threshold", 10),
        citation_reduction=getattr(request, "citation_severity_reduction", True),
    )

    if turnitin_stats["overall_score"] == 0 and any(r.sources for r in results):
        turnitin_stats = calculate_turnitin_score(
            request.text,
            results,
            exclude_quotes=True,
            exclude_bibliography=request.exclude_bibliography,
            small_match_threshold=getattr(request, "min_word_threshold", 10),
        )

    overall_score = turnitin_stats["overall_score"]
    plagiarism_percentage = turnitin_stats["overall_score"]
    plagiarized_count = sum(1 for r in results if r.is_plagiarized)

    used_ai_similarity = any(r.used_ai for r in results)
    fallback_used = any(r.fallback_used for r in results)
    analysis_method = (
        "hybrid"
        if used_ai_similarity and fallback_used
        else "semantic"
        if used_ai_similarity
        else "keyword"
    )

    filtered_results = results
    small_matches_removed = 0

    if request.exclude_small_matches > 0:
        filtered_results, small_matches_removed = _filter_small_matches(
            filtered_results, request.exclude_small_matches
        )
        exclusion_metadata["small_matches_removed"] = str(small_matches_removed)
        exclusion_metadata["small_match_threshold"] = str(request.exclude_small_matches)

    if request.exclude_small_sources:
        before_filter = len(filtered_results)
        filtered_results = _filter_small_sources(filtered_results, min_source_count=3)
        sources_removed = before_filter - len(filtered_results)
        exclusion_metadata["small_sources_removed"] = str(sources_removed)

    source_type_filter = getattr(request, "source_type_filter", None)
    if source_type_filter and len(source_type_filter) > 0:
        filtered_results = _filter_by_source_type(filtered_results, source_type_filter)
        exclusion_metadata["source_type_filter"] = ",".join(source_type_filter)

    contribution_threshold = getattr(request, "source_contribution_threshold", 0)
    if contribution_threshold > 0:
        filtered_results = _filter_by_contribution(
            filtered_results, contribution_threshold
        )
        exclusion_metadata["contribution_threshold"] = str(contribution_threshold)

    recalculate_for_response = len(filtered_results) > 0

    if recalculate_for_response:
        turnitin_stats_filtered = calculate_turnitin_score(
            request.text,
            filtered_results,
            exclude_quotes=True,
            exclude_bibliography=request.exclude_bibliography,
            small_match_threshold=10,
        )
        overall_score = turnitin_stats_filtered["overall_score"]
        plagiarism_percentage = turnitin_stats_filtered["overall_score"]
        plagiarized_count = sum(1 for r in filtered_results if r.is_plagiarized)
    else:
        overall_score = 0
        plagiarism_percentage = 0
        plagiarized_count = 0

    quota_info = get_quota_info(user_key=resolved_user_key)
    source_counts = (
        _build_source_counts(filtered_results) if recalculate_for_response else {}
    )

    if small_matches_removed > 0:
        exclusion_caveats.append(
            ReportV2Caveat(
                code="SMALL_MATCHES_FILTERED",
                message=f"{small_matches_removed} matches smaller than {request.exclude_small_matches} words were excluded from analysis.",
            )
        )

    ai_detection_score, ai_detection_confidence = _detect_ai_text(request.text)

    yield {
        "progress": 95,
        "current": total_sentences,
        "total": total_sentences,
        "status": "complete",
        "message": "Building final report...",
    }

    final_response = PlagiarismCheckResponse(
        overall_score=overall_score,
        plagiarism_percentage=plagiarism_percentage,
        total_sentences=len(filtered_results),
        plagiarized_sentences=plagiarized_count,
        results=filtered_results,
        used_ai_similarity=used_ai_similarity,
        fallback_used=fallback_used,
        analysis_method=analysis_method,
        ai_quota_remaining=quota_info["remaining"],
        ai_quota_percent=quota_info["usage_percent"],
        source_counts=source_counts,
        source_failures=None,
        quota_mode=quota_info.get("quota_mode"),
        ai_detection_score=ai_detection_score,
        ai_detection_confidence=ai_detection_confidence,
        report_v2=_build_report_v2(
            filtered_results,
            metadata_overrides=exclusion_metadata,
            extra_caveats=exclusion_caveats,
        ),
    )

    yield {
        "progress": 100,
        "current": total_sentences,
        "total": total_sentences,
        "status": "complete",
        "message": "Analysis complete",
    }

    yield final_response.model_dump()
